from __future__ import annotations

import base64
import binascii
import csv
import difflib
from email import policy
from email.message import EmailMessage
from email.parser import BytesParser
from email.utils import getaddresses
import json
import mimetypes
import re
import shutil
import subprocess
import tempfile
import zlib
import zipfile
from dataclasses import dataclass
from io import BytesIO
from io import StringIO
from hashlib import sha256
from pathlib import Path
import xml.etree.ElementTree as ET
from uuid import uuid4

from fastapi import APIRouter
from fastapi import Depends
from fastapi import File
from fastapi import Query
from fastapi import UploadFile
from fastapi.responses import FileResponse
from fastapi import Response
from PIL import Image
from PIL import ImageOps
from PIL import UnidentifiedImageError
from pillow_heif import register_heif_opener
from pydantic import BaseModel
from sqlalchemy import delete
from sqlalchemy import func
from sqlalchemy import or_
from sqlalchemy import select
from sqlalchemy.orm import Session as DatabaseSession

from caseclosed.auth import json_error
from caseclosed.db.models import Case
from caseclosed.db.models import Contact
from caseclosed.db.models import FileIconSetting
from caseclosed.db.models import FileLink
from caseclosed.db.models import FileSummary
from caseclosed.db.models import FileVersionDiff
from caseclosed.db.models import GmailMessage
from caseclosed.db.models import GmailMessageAttachment
from caseclosed.db.models import MailAutoState
from caseclosed.db.models import MailSummary
from caseclosed.db.models import MailThreadSummary
from caseclosed.db.models import MailUserState
from caseclosed.db.models import Paper
from caseclosed.db.models import PaperBibtexEntry
from caseclosed.db.models import PaperJournalIconSetting
from caseclosed.db.models import LlmRun
from caseclosed.db.models import StorageDirectory
from caseclosed.db.models import StorageLocation
from caseclosed.db.models import StorageObject
from caseclosed.db.models import StorageObjectVersion
from caseclosed.db.models import StoragePdfOcrStatus
from caseclosed.db.models import StorageOperationHistory
from caseclosed.db.models import Task
from caseclosed.db.runtime import BOOKSHELF_STORAGE_DIRECTORY_ID
from caseclosed.db.runtime import PAPERS_STORAGE_DIRECTORY_ID
from caseclosed.db.runtime import case_handover_storage_directory_id
from caseclosed.db.runtime import case_storage_directory_id
from caseclosed.db.runtime import get_session
from caseclosed.db.runtime import jst_iso
from caseclosed.settings import get_storage_root
from caseclosed.services.llm_provider import FUNCTION_TYPE_FILE_SUMMARY
from caseclosed.services.llm_provider import OpenAIProviderError
from caseclosed.services.llm_provider import build_file_summary_provider

router = APIRouter(prefix="/api/v1/storage", tags=["storage"])

register_heif_opener(thumbnails=False)

MAX_CONTACT_IMAGE_BYTES = 5 * 1024 * 1024
MAX_CONTACT_IMAGE_SIDE = 256
MAX_STORAGE_IMAGE_PREVIEW_SIDE = 2048
MAX_TMP_OBJECT_BYTES = 1024 * 1024 * 1024
MAX_ARCHIVE_TREE_ENTRIES = 1000
MAX_EML_PREVIEW_BYTES = 25 * 1024 * 1024
MAX_FILE_SUMMARY_INPUT_BYTES = 900 * 1024
MAX_FILE_SUMMARY_INPUT_CHARS = 60000
MAX_PAPER_BIBTEX_BYTES = 2 * 1024 * 1024
CONTACT_IMAGE_CONTENT_TYPES = {
    "image/png",
    "image/jpeg",
    "image/gif",
    "image/heic",
    "image/heif",
    "image/webp",
    "image/svg+xml",
}
FILE_ICON_CONTENT_TYPES = CONTACT_IMAGE_CONTENT_TYPES
MAX_FILE_ICON_BYTES = 5 * 1024 * 1024
CONTENT_TYPE_EXTENSIONS = {
    "image/png": ".png",
    "image/jpeg": ".jpg",
    "image/gif": ".gif",
    "image/heic": ".heic",
    "image/heif": ".heif",
    "image/webp": ".webp",
    "image/svg+xml": ".svg",
    "application/pdf": ".pdf",
    "application/zip": ".zip",
}
INTERNAL_STORAGE_LOCATION_ID = "storage_location_internal"
GMAIL_ATTACHMENT_STORAGE_SCOPE = "tmp/gmail-attachments"
NO_STORE_FILE_HEADERS = {"Cache-Control": "no-store"}
CACHEABLE_IMAGE_FILE_HEADERS = {"Cache-Control": "private, max-age=604800, immutable"}
ACTIVE_STORAGE_DIRECTORY_KINDS = ("normal", "case", "task")
HIDDEN_STORAGE_DIRECTORY_IDS = {BOOKSHELF_STORAGE_DIRECTORY_ID, PAPERS_STORAGE_DIRECTORY_ID}


class ContactImageUpload(BaseModel):
    filename: str | None = None
    content_type: str
    data_base64: str


class TemporaryObjectUpload(BaseModel):
    filename: str | None = None
    content_type: str | None = None
    data_base64: str


class ManagedObjectUpload(BaseModel):
    filename: str | None = None
    content_type: str | None = None
    data_base64: str
    directory_id: str | None = None


class FileIconSettingCreate(BaseModel):
    icon_filename: str | None = None
    icon_content_type: str
    icon_data_base64: str
    extensions: list[str]


class FileIconSettingUpdate(BaseModel):
    icon_filename: str | None = None
    icon_content_type: str | None = None
    icon_data_base64: str | None = None
    extensions: list[str] | None = None


class StorageObjectLlmInputPatch(BaseModel):
    llm_input_allowed: bool


class StorageObjectFilenamePatch(BaseModel):
    filename: str | None = None


class StorageObjectDirectoryPatch(BaseModel):
    directory_id: str | None = None


class StorageDirectoryParentPatch(BaseModel):
    parent_id: str | None = None


class StorageObjectCaseLinkCreate(BaseModel):
    case_id: str


class StorageDirectoryCreate(BaseModel):
    name: str
    parent_id: str | None = None


class FileSummaryRequest(BaseModel):
    storage_object_version_id: str | None = None


class PaperUpdate(BaseModel):
    title: str | None = None
    authors_text: str | None = None
    bibtex: str | None = None
    summary: str | None = None
    tags: list[str] | None = None


class PaperJournalIconSettingCreate(BaseModel):
    match_journal: str
    icon_filename: str | None = None
    icon_content_type: str
    icon_data_base64: str


class PaperJournalIconSettingUpdate(BaseModel):
    match_journal: str | None = None
    icon_filename: str | None = None
    icon_content_type: str | None = None
    icon_data_base64: str | None = None


@dataclass(frozen=True)
class FileTextExtraction:
    source_kind: str
    read_scope: str
    source_text: str
    truncated: bool
    limitations: list[str]


def format_bytes(value: int) -> str:
    if value < 1024:
        return f"{value} B"
    kib = value / 1024
    if kib < 1024:
        return f"{kib:.0f} KB" if kib >= 10 else f"{kib:.1f} KB"
    mib = kib / 1024
    if mib < 1024:
        return f"{mib:.0f} MB" if mib >= 10 else f"{mib:.1f} MB"
    gib = mib / 1024
    return f"{gib:.0f} GB" if gib >= 10 else f"{gib:.1f} GB"


def new_id(prefix: str) -> str:
    return f"{prefix}_{uuid4().hex}"


def normalize_file_icon_extensions(values: list[str]) -> list[str]:
    extensions: list[str] = []
    seen: set[str] = set()
    for raw_value in values:
        for part in re.split(r"[,\s]+", raw_value):
            value = part.strip().lower()
            if value == "":
                continue
            if value.startswith("*."):
                value = value[1:]
            if not value.startswith("."):
                value = f".{value}"
            if not re.fullmatch(r"\.[a-z0-9][a-z0-9_-]{0,31}", value):
                raise json_error(422, "VALIDATION_ERROR", "Invalid file extension.")
            if value not in seen:
                seen.add(value)
                extensions.append(value)
    if len(extensions) == 0:
        raise json_error(422, "VALIDATION_ERROR", "At least one extension is required.")
    return extensions




def is_pdf_filename(filename: str | None) -> bool:
    return filename is not None and filename.strip().lower().endswith(".pdf")


def is_citation_filename(filename: str | None) -> bool:
    if filename is None:
        return False
    lowered = filename.strip().lower()
    return lowered.endswith(".bib") or lowered.endswith(".bibtex") or lowered.endswith(".ris") or lowered.endswith(".nbib")


def is_bookshelf_pdf_filename(filename: str | None) -> bool:
    return is_pdf_filename(filename)


def is_bookshelf_epub_filename(filename: str | None) -> bool:
    return filename is not None and filename.strip().lower().endswith(".epub")


def is_bookshelf_material_filename(filename: str | None) -> bool:
    return is_bookshelf_pdf_filename(filename) or is_bookshelf_epub_filename(filename)


def converted_bookshelf_pdf_filename(filename: str | None) -> str:
    raw_name = (filename or "book.epub").strip() or "book.epub"
    return f"{Path(raw_name).stem or 'book'}.pdf"


def normalized_optional_filename(value: str | None) -> str | None:
    if value is None:
        return None
    stripped = value.strip()
    return stripped if stripped != "" else None


def file_icon_setting_data(setting: FileIconSetting) -> dict[str, object]:
    try:
        extensions = json.loads(setting.extensions_json)
    except json.JSONDecodeError:
        extensions = []
    if not isinstance(extensions, list):
        extensions = []
    icon_url = (
        storage_object_url(setting.storage_object_id)
        if setting.storage_object_id is not None
        else setting.icon_data_url
    )
    return {
        "id": setting.id,
        "storage_object_id": setting.storage_object_id,
        "icon_filename": setting.icon_filename,
        "icon_content_type": setting.icon_content_type,
        "icon_url": icon_url,
        "icon_data_url": setting.icon_data_url,
        "extensions": [str(extension) for extension in extensions],
        "created_at": setting.created_at,
        "updated_at": setting.updated_at,
        "version": setting.version,
    }


def prepare_file_icon_image(*, content_type: str, data_base64: str) -> tuple[str, bytes]:
    content_type = content_type.strip().lower()
    if content_type not in FILE_ICON_CONTENT_TYPES:
        raise json_error(422, "VALIDATION_ERROR", "Unsupported file icon image type.")
    data = decode_base64_payload(data_base64)
    if len(data) == 0:
        raise json_error(422, "VALIDATION_ERROR", "File icon image is empty.")
    if len(data) > MAX_FILE_ICON_BYTES:
        raise json_error(413, "PAYLOAD_TOO_LARGE", "File icon image is too large.")
    return resize_contact_image(content_type=content_type, data=data)


def file_icon_for_filename(
    session: DatabaseSession,
    filename: str | None,
) -> FileIconSetting | None:
    extension = normalized_file_extension(filename)
    if extension == "":
        return None
    settings = session.scalars(
        select(FileIconSetting).order_by(
            FileIconSetting.created_at.asc(),
            FileIconSetting.id.asc(),
        )
    ).all()
    for setting in settings:
        try:
            extensions = json.loads(setting.extensions_json)
        except json.JSONDecodeError:
            continue
        if isinstance(extensions, list) and extension in {str(item) for item in extensions}:
            return setting
    return None


def storage_source_mail_data(
    session: DatabaseSession,
    message_id: str | None,
) -> dict[str, object] | None:
    if message_id is None:
        return None
    message = session.get(GmailMessage, message_id)
    if message is None:
        return None
    user_state = session.scalar(
        select(MailUserState).where(MailUserState.message_id == message.id)
    )
    auto_state = session.scalar(
        select(MailAutoState).where(MailAutoState.message_id == message.id)
    )
    effective_importance = (
        user_state.user_importance
        if user_state is not None and user_state.user_importance is not None
        else auto_state.effective_importance if auto_state is not None else "pending"
    )
    summary = session.scalar(
        select(MailThreadSummary).where(MailThreadSummary.thread_id == message.thread_id)
    )
    if summary is None:
        summary = session.scalar(
            select(MailSummary)
            .where(MailSummary.message_id == message.id)
            .limit(1)
        )
    return {
        "id": message.id,
        "received_at": message.received_at,
        "subject": message.subject,
        "from_address": message.from_address,
        "from_name": message.from_name,
        "effective_importance": effective_importance,
        "read_status": user_state.read_status if user_state is not None else "read",
        "summary": summary.summary_text if summary is not None else message.snippet,
        "has_attachments": True,
    }


def storage_object_data(
    storage_object: StorageObject,
    session: DatabaseSession | None = None,
    *,
    display_source: str | None = None,
) -> dict[str, object]:
    data = {
        "id": storage_object.id,
        "directory_id": storage_object.directory_id,
        "location_id": storage_object.location_id,
        "scope": storage_object.scope,
        "original_filename": storage_object.original_filename,
        "content_type": storage_object.content_type,
        "byte_size": storage_object.byte_size,
        "sha256_hex": storage_object.sha256_hex,
        "llm_input_allowed": bool(storage_object.llm_input_allowed),
        "source_type": storage_object.source_type,
        "source_message_id": storage_object.source_message_id,
        "url": storage_object_url(storage_object.id),
        "created_at": storage_object.created_at,
        "updated_at": storage_object.updated_at,
        "file_updated_at": storage_object.file_updated_at,
    }
    if display_source is not None:
        data["display_source"] = display_source
    if session is not None:
        data["source_mail"] = storage_source_mail_data(
            session,
            storage_object.source_message_id,
        )
        data["directory_path"] = storage_directory_path(session, storage_object.directory_id)
        data["physical_directory_id"] = storage_object.directory_id
        data["physical_directory_path"] = data["directory_path"]
        storage_directory_param = storage_object.directory_id or "root"
        data["physical_directory_url"] = f"/storage?directory={storage_directory_param}"
        file_icon = file_icon_for_filename(session, storage_object.original_filename)
        data["file_icon_setting_id"] = file_icon.id if file_icon is not None else None
        data["file_icon_url"] = (
            storage_object_url(file_icon.storage_object_id)
            if file_icon is not None and file_icon.storage_object_id is not None
            else file_icon.icon_data_url if file_icon is not None else None
        )
    return data


def storage_object_version_data(version: StorageObjectVersion) -> dict[str, object]:
    return {
        "id": version.id,
        "storage_object_id": version.storage_object_id,
        "version_number": version.version_number,
        "original_filename": version.original_filename,
        "content_type": version.content_type,
        "byte_size": version.byte_size,
        "sha256_hex": version.sha256_hex,
        "url": (
            f"/api/v1/storage/objects/{version.storage_object_id}"
            f"/versions/{version.id}/content"
        ),
        "download_url": (
            f"/api/v1/storage/objects/{version.storage_object_id}"
            f"/versions/{version.id}/download"
        ),
        "created_at": version.created_at,
    }


def safe_json_list(value: str | None) -> list[object]:
    if value is None:
        return []
    try:
        parsed = json.loads(value)
    except json.JSONDecodeError:
        return []
    return parsed if isinstance(parsed, list) else []


def safe_json_dict(value: str | None) -> dict[str, object]:
    if value is None:
        return {}
    try:
        parsed = json.loads(value)
    except json.JSONDecodeError:
        return {}
    return parsed if isinstance(parsed, dict) else {}


def file_summary_data(summary: FileSummary) -> dict[str, object]:
    return {
        "id": summary.id,
        "storage_object_id": summary.storage_object_id,
        "storage_object_version_id": summary.storage_object_version_id,
        "source_sha256_hex": summary.source_sha256_hex,
        "source_filename": summary.source_filename,
        "source_content_type": summary.source_content_type,
        "source_byte_size": summary.source_byte_size,
        "summary_type": summary.summary_type,
        "file_description": summary.file_description,
        "summary_points": safe_json_list(summary.summary_points_json)[:5],
        "llm_digest": summary.llm_digest,
        "structured_digest": safe_json_dict(summary.structured_digest_json),
        "coverage": safe_json_dict(summary.coverage_json),
        "token_estimate": summary.token_estimate,
        "llm_run_id": summary.llm_run_id,
        "created_at": summary.created_at,
        "updated_at": summary.updated_at,
        "version": summary.version,
    }


def file_version_diff_data(
    diff: FileVersionDiff,
    *,
    display_lines: list[dict[str, str]] | None = None,
) -> dict[str, object]:
    return {
        "id": diff.id,
        "storage_object_id": diff.storage_object_id,
        "previous_version_id": diff.previous_version_id,
        "previous_sha256_hex": diff.previous_sha256_hex,
        "current_sha256_hex": diff.current_sha256_hex,
        "diff_kind": diff.diff_kind,
        "summary_text": diff.summary_text,
        "added_lines": safe_json_list(diff.added_lines_json),
        "removed_lines": safe_json_list(diff.removed_lines_json),
        "display_lines": display_lines or [],
        "coverage": safe_json_dict(diff.coverage_json),
        "created_at": diff.created_at,
        "updated_at": diff.updated_at,
        "version": diff.version,
    }


def storage_directory_is_hidden_from_storage_ui(directory_id: str | None) -> bool:
    return directory_id in HIDDEN_STORAGE_DIRECTORY_IDS


def storage_directory_data(directory: StorageDirectory) -> dict[str, object]:
    return {
        "id": directory.id,
        "parent_id": directory.parent_id,
        "directory_kind": directory.directory_kind,
        "case_id": directory.case_id,
        "name": directory.name,
        "status": directory.status,
        "created_at": directory.created_at,
        "updated_at": directory.updated_at,
        "version": directory.version,
    }


def normalize_directory_id(value: str | None) -> str | None:
    if value is None:
        return None
    stripped = value.strip()
    return None if stripped in {"", "root", "all"} else stripped


def ensure_storage_directory(
    session: DatabaseSession,
    directory_id: str | None,
) -> StorageDirectory | None:
    normalized_id = normalize_directory_id(directory_id)
    if normalized_id is None:
        return None
    directory = session.get(StorageDirectory, normalized_id)
    if (
        directory is None
        or directory.status != "active"
        or directory.directory_kind not in ACTIVE_STORAGE_DIRECTORY_KINDS
    ):
        raise json_error(404, "NOT_FOUND", "Storage directory not found.")
    return directory


def storage_directory_breadcrumbs(
    session: DatabaseSession,
    directory_id: str | None,
) -> list[dict[str, object]]:
    breadcrumbs: list[dict[str, object]] = []
    seen: set[str] = set()
    current = ensure_storage_directory(session, directory_id)
    while current is not None:
        if current.id in seen:
            break
        seen.add(current.id)
        breadcrumbs.append(storage_directory_data(current))
        current = ensure_storage_directory(session, current.parent_id)
    breadcrumbs.reverse()
    return breadcrumbs


def storage_directory_path(
    session: DatabaseSession,
    directory_id: str | None,
) -> list[str]:
    return [
        str(item["name"])
        for item in storage_directory_breadcrumbs(session, directory_id)
    ]


def storage_directory_descendant_ids(
    session: DatabaseSession,
    directory_id: str | None,
) -> set[str]:
    directories = session.scalars(
        select(StorageDirectory).where(
            StorageDirectory.status == "active",
            StorageDirectory.directory_kind.in_(ACTIVE_STORAGE_DIRECTORY_KINDS),
            StorageDirectory.id.not_in(HIDDEN_STORAGE_DIRECTORY_IDS),
        )
    ).all()
    if directory_id is None:
        return {directory.id for directory in directories}
    children: dict[str | None, list[str]] = {}
    for directory in directories:
        children.setdefault(directory.parent_id, []).append(directory.id)
    descendant_ids: set[str] = set()
    stack = [directory_id]
    while stack:
        current_id = stack.pop()
        if current_id in descendant_ids:
            continue
        descendant_ids.add(current_id)
        stack.extend(children.get(current_id, []))
    return descendant_ids


def storage_directory_case_id(
    session: DatabaseSession,
    directory_id: str | None,
) -> str | None:
    current = ensure_storage_directory(session, directory_id)
    seen: set[str] = set()
    while current is not None:
        if current.id in seen:
            return None
        seen.add(current.id)
        if current.directory_kind == "case" and current.case_id is not None:
            return current.case_id
        current = ensure_storage_directory(session, current.parent_id)
    return None


def is_case_handover_directory_descendant(
    session: DatabaseSession,
    directory_id: str | None,
) -> bool:
    current = ensure_storage_directory(session, directory_id)
    seen: set[str] = set()
    while current is not None:
        if current.id in seen:
            return False
        seen.add(current.id)
        if (
            current.case_id is not None
            and current.id == case_handover_storage_directory_id(current.case_id)
        ):
            return True
        current = ensure_storage_directory(session, current.parent_id)
    return False


def storage_object_linked_case_data(
    case: Case,
    *,
    source: str,
    file_link: FileLink | None = None,
) -> dict[str, object]:
    return {
        "case_id": case.id,
        "case_name": case.name,
        "source": source,
        "file_link_id": file_link.id if file_link is not None else None,
        "created_at": file_link.created_at if file_link is not None else None,
        "updated_at": file_link.updated_at if file_link is not None else case.updated_at,
        "case_url": f"/cases/{case.id}",
    }


def storage_object_linked_case_items(
    session: DatabaseSession,
    storage_object: StorageObject,
) -> list[dict[str, object]]:
    items_by_case_id: dict[str, dict[str, object]] = {}
    physical_case_id = storage_directory_case_id(session, storage_object.directory_id)
    if physical_case_id is not None:
        physical_case = session.get(Case, physical_case_id)
        if physical_case is not None:
            items_by_case_id[physical_case.id] = storage_object_linked_case_data(
                physical_case,
                source="physical",
            )

    rows = session.execute(
        select(FileLink, Case)
        .join(Case, Case.id == FileLink.linked_id)
        .where(FileLink.storage_object_id == storage_object.id)
        .where(FileLink.linked_type == "case")
        .where(FileLink.status == "active")
        .order_by(FileLink.created_at.desc(), FileLink.id.desc())
    ).all()
    for file_link, case in rows:
        if case.id in items_by_case_id:
            continue
        items_by_case_id[case.id] = storage_object_linked_case_data(
            case,
            source="link",
            file_link=file_link,
        )

    return sorted(
        items_by_case_id.values(),
        key=lambda item: (str(item["case_name"]).lower(), str(item["case_id"])),
    )


def storage_case_for_directory_id(
    session: DatabaseSession,
    directory_id: str | None,
) -> Case | None:
    if directory_id is None:
        return None
    case_id = storage_directory_case_id(session, directory_id)
    if case_id is None:
        return None
    return session.get(Case, case_id)


def storage_linked_case_directory_objects(
    session: DatabaseSession,
    case: Case,
    directory_id: str,
    *,
    status: str,
    location_id: str,
    exclude_ids: set[str],
    limit: int,
) -> list[StorageObject]:
    if limit <= 0:
        return []
    linked_directory_conditions = [FileLink.directory_id == directory_id]
    if directory_id == case_storage_directory_id(case.id):
        linked_directory_conditions.append(FileLink.directory_id.is_(None))
    statement = (
        select(StorageObject)
        .join(FileLink, FileLink.storage_object_id == StorageObject.id)
        .where(FileLink.linked_type == "case")
        .where(FileLink.linked_id == case.id)
        .where(FileLink.status == "active")
        .where(or_(*linked_directory_conditions))
        .where(StorageObject.scope == "managed")
        .where(StorageObject.id.not_in(exclude_ids))
        .order_by(StorageObject.created_at.desc(), StorageObject.id.desc())
        .limit(limit)
    )
    if status != "all":
        statement = statement.where(StorageObject.status == status)
    if location_id != "all":
        statement = statement.where(StorageObject.location_id == location_id)
    return session.scalars(statement).all()


def sort_storage_objects_for_directory(items: list[StorageObject]) -> list[StorageObject]:
    return sorted(
        items,
        key=lambda item: (item.created_at, item.id),
        reverse=True,
    )


def storage_file_extension(filename: str | None) -> str | None:
    if filename is None or "." not in filename:
        return None
    extension = filename.rsplit(".", 1)[1].strip().lower()
    return extension or None


def storage_location_data(
    location: StorageLocation,
    *,
    object_count: int = 0,
    active_byte_size: int = 0,
) -> dict[str, object]:
    return {
        "id": location.id,
        "label": location.label,
        "kind": location.kind,
        "root_path": location.root_path,
        "mount_hint": location.mount_hint,
        "marker_id": location.marker_id,
        "status": location.status,
        "object_count": object_count,
        "active_byte_size": active_byte_size,
        "created_at": location.created_at,
        "updated_at": location.updated_at,
        "version": location.version,
    }


def storage_object_url(storage_object_id: str) -> str:
    return f"/api/v1/storage/objects/{storage_object_id}/content"


def should_record_storage_content_view(storage_object: StorageObject) -> bool:
    return storage_object.scope == "managed"


def storage_content_response_headers(storage_object: StorageObject) -> dict[str, str]:
    if (
        storage_object.scope != "managed"
        and storage_object.content_type.lower().startswith("image/")
    ):
        return CACHEABLE_IMAGE_FILE_HEADERS
    return NO_STORE_FILE_HEADERS


def storage_object_id_from_url(value: str | None) -> str | None:
    if value is None:
        return None
    prefix = "/api/v1/storage/objects/"
    suffix = "/content"
    if not value.startswith(prefix) or not value.endswith(suffix):
        return None
    storage_object_id = value.removeprefix(prefix).removesuffix(suffix)
    if storage_object_id == "":
        return None
    return storage_object_id


def storage_root() -> Path:
    root = get_storage_root()
    root.mkdir(parents=True, exist_ok=True)
    (root / "tmp").mkdir(parents=True, exist_ok=True)
    (root / "contact-images").mkdir(parents=True, exist_ok=True)
    (root / "file-icons").mkdir(parents=True, exist_ok=True)
    (root / "case-tool-icons").mkdir(parents=True, exist_ok=True)
    return root


def storage_location_root(
    session: DatabaseSession,
    location_id: str | None,
) -> Path:
    effective_location_id = location_id or INTERNAL_STORAGE_LOCATION_ID
    location = session.get(StorageLocation, effective_location_id)
    if location is None or location.status != "active":
        raise RuntimeError(f"Storage location is not available: {effective_location_id}")
    root = Path(location.root_path)
    root.mkdir(parents=True, exist_ok=True)
    return root


def storage_object_absolute_path(
    storage_object: StorageObject,
    session: DatabaseSession,
) -> Path:
    root = storage_location_root(session, storage_object.location_id).resolve()
    object_path = (root / storage_object.storage_path).resolve()
    if not object_path.is_relative_to(root):
        raise RuntimeError("Storage object path escapes storage root.")
    return object_path


def storage_object_version_absolute_path(
    storage_object: StorageObject,
    version: StorageObjectVersion,
    session: DatabaseSession,
) -> Path:
    root = storage_location_root(session, storage_object.location_id).resolve()
    version_path = (root / version.storage_path).resolve()
    if not version_path.is_relative_to(root):
        raise RuntimeError("Storage object version path escapes storage root.")
    return version_path


def decode_base64_payload(value: str) -> bytes:
    raw_value = value.strip()
    if "," in raw_value and raw_value.lower().startswith("data:"):
        raw_value = raw_value.split(",", maxsplit=1)[1]
    try:
        return base64.b64decode(raw_value, validate=True)
    except binascii.Error as error:
        raise json_error(422, "VALIDATION_ERROR", "Invalid base64 data.") from error


def resize_contact_image(
    *,
    content_type: str,
    data: bytes,
) -> tuple[str, bytes]:
    if content_type == "image/svg+xml":
        return content_type, data

    try:
        with Image.open(BytesIO(data)) as image:
            image = ImageOps.exif_transpose(image)
            image.thumbnail(
                (MAX_CONTACT_IMAGE_SIDE, MAX_CONTACT_IMAGE_SIDE),
                Image.Resampling.LANCZOS,
            )
            output = BytesIO()
            if image.mode not in {"RGB", "RGBA"}:
                image = image.convert("RGBA")
            image.save(output, format="WEBP", quality=85, method=6)
            return "image/webp", output.getvalue()
    except (UnidentifiedImageError, OSError) as error:
        raise json_error(422, "VALIDATION_ERROR", "Invalid contact image data.") from error


def is_heic_image(*, content_type: str | None, filename: str | None) -> bool:
    normalized_content_type = (content_type or "").lower().split(";", maxsplit=1)[0].strip()
    if normalized_content_type in {
        "image/heic",
        "image/heic-sequence",
        "image/heif",
        "image/heif-sequence",
    }:
        return True
    return Path(filename or "").suffix.lower() in {".heic", ".heif"}


def storage_image_preview_response(
    *,
    source_path: Path,
    content_type: str | None,
    filename: str | None,
    sha256_hex: str,
) -> Response:
    if not is_heic_image(content_type=content_type, filename=filename):
        raise json_error(415, "UNSUPPORTED_MEDIA_TYPE", "Image preview conversion is not required.")
    try:
        with Image.open(source_path) as source_image:
            image = ImageOps.exif_transpose(source_image)
            image.thumbnail(
                (MAX_STORAGE_IMAGE_PREVIEW_SIDE, MAX_STORAGE_IMAGE_PREVIEW_SIDE),
                Image.Resampling.LANCZOS,
            )
            if image.mode not in {"RGB", "RGBA"}:
                image = image.convert("RGBA")
            output = BytesIO()
            image.save(output, format="WEBP", quality=85, method=6)
    except (UnidentifiedImageError, OSError, ValueError) as error:
        raise json_error(422, "INVALID_IMAGE", "Storage image preview could not be generated.") from error
    return Response(
        content=output.getvalue(),
        media_type="image/webp",
        headers={
            **CACHEABLE_IMAGE_FILE_HEADERS,
            "ETag": f'"{sha256_hex}-storage-preview-v1"',
        },
    )


def object_relative_path(
    *,
    scope: str,
    storage_object_id: str,
    content_type: str | None,
) -> Path:
    extension = CONTENT_TYPE_EXTENSIONS.get(content_type or "", ".bin")
    bucket = storage_object_id.removeprefix("storage_object_")[:2] or "00"
    return Path(scope) / bucket / f"{storage_object_id}{extension}"


def save_storage_object(
    session: DatabaseSession,
    *,
    scope: str,
    filename: str | None,
    content_type: str | None,
    data: bytes,
    now: str,
    directory_id: str | None = None,
    source_type: str | None = None,
    source_message_id: str | None = None,
) -> StorageObject:
    object_id = new_id("storage_object")
    relative_path = object_relative_path(
        scope=scope,
        storage_object_id=object_id,
        content_type=content_type,
    )
    root = storage_root()
    absolute_path = root / relative_path
    absolute_path.parent.mkdir(parents=True, exist_ok=True)
    absolute_path.write_bytes(data)
    storage_object = StorageObject(
        id=object_id,
        directory_id=directory_id,
        location_id=INTERNAL_STORAGE_LOCATION_ID,
        scope=scope,
        original_filename=filename,
        content_type=content_type,
        byte_size=len(data),
        sha256_hex=sha256(data).hexdigest(),
        storage_path=relative_path.as_posix(),
        llm_input_allowed=1 if is_case_handover_directory_descendant(session, directory_id) else 0,
        source_type=source_type,
        source_message_id=source_message_id,
        status="active",
        created_at=now,
        updated_at=now,
        file_updated_at=now,
        version=1,
    )
    session.add(storage_object)
    record_storage_operation(
        session,
        operation_type="created",
        now=now,
        storage_object=storage_object,
        details={"scope": scope, "source_type": source_type},
    )
    return storage_object


async def save_uploaded_storage_object(
    session: DatabaseSession,
    *,
    scope: str,
    upload: UploadFile,
    now: str,
    directory_id: str | None = None,
    source_type: str | None = None,
    source_message_id: str | None = None,
) -> StorageObject:
    content_type = (upload.content_type or "application/octet-stream").strip()
    if content_type == "":
        content_type = "application/octet-stream"
    object_id = new_id("storage_object")
    relative_path = object_relative_path(
        scope=scope,
        storage_object_id=object_id,
        content_type=content_type,
    )
    root = storage_root()
    absolute_path = root / relative_path
    absolute_path.parent.mkdir(parents=True, exist_ok=True)

    hasher = sha256()
    byte_size = 0
    try:
        with absolute_path.open("wb") as handle:
            while True:
                chunk = await upload.read(1024 * 1024)
                if chunk == b"":
                    break
                byte_size += len(chunk)
                if byte_size > MAX_TMP_OBJECT_BYTES:
                    raise json_error(
                        413,
                        "PAYLOAD_TOO_LARGE",
                        "Storage object is too large.",
                    )
                hasher.update(chunk)
                handle.write(chunk)
    except Exception:
        if absolute_path.is_file():
            absolute_path.unlink()
        raise
    finally:
        await upload.close()

    if byte_size == 0:
        if absolute_path.is_file():
            absolute_path.unlink()
        raise json_error(422, "VALIDATION_ERROR", "Storage object is empty.")

    storage_object = StorageObject(
        id=object_id,
        directory_id=directory_id,
        location_id=INTERNAL_STORAGE_LOCATION_ID,
        scope=scope,
        original_filename=upload.filename,
        content_type=content_type,
        byte_size=byte_size,
        sha256_hex=hasher.hexdigest(),
        storage_path=relative_path.as_posix(),
        llm_input_allowed=1 if is_case_handover_directory_descendant(session, directory_id) else 0,
        source_type=source_type,
        source_message_id=source_message_id,
        status="active",
        created_at=now,
        updated_at=now,
        file_updated_at=now,
        version=1,
    )
    session.add(storage_object)
    record_storage_operation(
        session,
        operation_type="uploaded",
        now=now,
        storage_object=storage_object,
        details={"scope": scope, "source_type": source_type},
    )
    return storage_object


def storage_object_version_relative_path(
    storage_object: StorageObject,
    *,
    version_id: str,
    version_number: int,
) -> Path:
    suffix = Path(storage_object.storage_path).suffix or ".bin"
    return (
        Path(storage_object.scope)
        / "versions"
        / storage_object.id
        / f"v{version_number}_{version_id}{suffix}"
    )


def next_storage_object_version_number(
    session: DatabaseSession,
    storage_object_id: str,
) -> int:
    current = session.scalar(
        select(func.max(StorageObjectVersion.version_number)).where(
            StorageObjectVersion.storage_object_id == storage_object_id
        )
    )
    return int(current or 0) + 1


async def update_storage_object_from_upload(
    session: DatabaseSession,
    storage_object: StorageObject,
    *,
    upload: UploadFile,
    now: str,
) -> StorageObjectVersion | None:
    current_path = storage_object_absolute_path(storage_object, session)
    if not current_path.is_file():
        raise json_error(404, "NOT_FOUND", "Storage object file not found.")

    content_type = (upload.content_type or "application/octet-stream").strip()
    if content_type == "":
        content_type = "application/octet-stream"
    new_relative_path = object_relative_path(
        scope=storage_object.scope,
        storage_object_id=storage_object.id,
        content_type=content_type,
    )
    root = storage_location_root(session, storage_object.location_id).resolve()
    new_path = (root / new_relative_path).resolve()
    if not new_path.is_relative_to(root):
        raise RuntimeError("Storage object path escapes storage root.")

    incoming_path = (
        root
        / storage_object.scope
        / "_incoming"
        / f"{new_id('storage_upload')}.tmp"
    ).resolve()
    if not incoming_path.is_relative_to(root):
        raise RuntimeError("Storage upload path escapes storage root.")
    incoming_path.parent.mkdir(parents=True, exist_ok=True)

    hasher = sha256()
    byte_size = 0
    try:
        with incoming_path.open("wb") as handle:
            while True:
                chunk = await upload.read(1024 * 1024)
                if chunk == b"":
                    break
                byte_size += len(chunk)
                if byte_size > MAX_TMP_OBJECT_BYTES:
                    raise json_error(
                        413,
                        "PAYLOAD_TOO_LARGE",
                        "Storage object is too large.",
                    )
                hasher.update(chunk)
                handle.write(chunk)
    except Exception:
        if incoming_path.is_file():
            incoming_path.unlink()
        raise
    finally:
        await upload.close()

    if byte_size == 0:
        if incoming_path.is_file():
            incoming_path.unlink()
        raise json_error(422, "VALIDATION_ERROR", "Storage object is empty.")

    new_sha256_hex = hasher.hexdigest()
    if byte_size == storage_object.byte_size and new_sha256_hex == storage_object.sha256_hex:
        if incoming_path.is_file():
            incoming_path.unlink()
        record_storage_operation(
            session,
            operation_type="update_skipped",
            now=now,
            storage_object=storage_object,
            details={
                "reason": "duplicate_content",
                "incoming_filename": upload.filename,
                "incoming_content_type": content_type,
                "incoming_byte_size": byte_size,
            },
        )
        return None

    version_number = next_storage_object_version_number(session, storage_object.id)
    version_id = new_id("storage_object_version")
    version_relative_path = storage_object_version_relative_path(
        storage_object,
        version_id=version_id,
        version_number=version_number,
    )
    version_path = (root / version_relative_path).resolve()
    if not version_path.is_relative_to(root):
        raise RuntimeError("Storage object version path escapes storage root.")

    previous_original_filename = storage_object.original_filename
    previous_content_type = storage_object.content_type
    previous_byte_size = storage_object.byte_size
    previous_sha256_hex = storage_object.sha256_hex
    previous_storage_path = storage_object.storage_path
    previous_file_updated_at = storage_object.file_updated_at

    version_path.parent.mkdir(parents=True, exist_ok=True)
    moved_current = False
    try:
        current_path.replace(version_path)
        moved_current = True
        new_path.parent.mkdir(parents=True, exist_ok=True)
        if new_path.is_file():
            new_path.unlink()
        incoming_path.replace(new_path)
    except Exception:
        if incoming_path.is_file():
            incoming_path.unlink()
        if moved_current and version_path.is_file() and not current_path.exists():
            version_path.replace(current_path)
        raise

    version = StorageObjectVersion(
        id=version_id,
        storage_object_id=storage_object.id,
        version_number=version_number,
        original_filename=previous_original_filename,
        content_type=previous_content_type,
        byte_size=previous_byte_size,
        sha256_hex=previous_sha256_hex,
        storage_path=version_relative_path.as_posix(),
        created_at=previous_file_updated_at,
    )
    session.add(version)

    storage_object.original_filename = upload.filename
    storage_object.content_type = content_type
    storage_object.byte_size = byte_size
    storage_object.sha256_hex = new_sha256_hex
    storage_object.storage_path = new_relative_path.as_posix()
    storage_object.updated_at = now
    storage_object.file_updated_at = now
    storage_object.version += 1
    reassign_current_file_summaries_to_version(
        session,
        storage_object=storage_object,
        previous_version=version,
        previous_sha256_hex=previous_sha256_hex,
        now=now,
    )
    prepare_file_version_diff(
        session,
        storage_object=storage_object,
        previous_version=version,
        now=now,
    )
    record_storage_operation(
        session,
        operation_type="updated",
        now=now,
        storage_object=storage_object,
        details={
            "previous_version_id": version.id,
            "previous_version_number": version.version_number,
            "previous_storage_path": previous_storage_path,
            "previous_file_updated_at": previous_file_updated_at,
            "new_filename": upload.filename,
        },
    )
    return version


def update_storage_object_from_bytes(
    session: DatabaseSession,
    storage_object: StorageObject,
    *,
    data: bytes,
    filename: str | None,
    content_type: str,
    now: str,
    operation_type: str = "updated",
    operation_details: dict[str, object] | None = None,
) -> StorageObjectVersion | None:
    if len(data) == 0:
        raise json_error(422, "VALIDATION_ERROR", "Storage object is empty.")
    current_path = storage_object_absolute_path(storage_object, session)
    if not current_path.is_file():
        raise json_error(404, "NOT_FOUND", "Storage object file not found.")

    new_sha256_hex = sha256(data).hexdigest()
    if len(data) == storage_object.byte_size and new_sha256_hex == storage_object.sha256_hex:
        record_storage_operation(
            session,
            operation_type="update_skipped",
            now=now,
            storage_object=storage_object,
            details={
                "reason": "duplicate_content",
                "incoming_filename": filename,
                "incoming_content_type": content_type,
                "incoming_byte_size": len(data),
            },
        )
        return None

    root = storage_location_root(session, storage_object.location_id).resolve()
    new_relative_path = object_relative_path(
        scope=storage_object.scope,
        storage_object_id=storage_object.id,
        content_type=content_type,
    )
    new_path = (root / new_relative_path).resolve()
    if not new_path.is_relative_to(root):
        raise RuntimeError("Storage object path escapes storage root.")

    incoming_path = (
        root
        / storage_object.scope
        / "_incoming"
        / f"{new_id('storage_upload')}.tmp"
    ).resolve()
    if not incoming_path.is_relative_to(root):
        raise RuntimeError("Storage upload path escapes storage root.")
    incoming_path.parent.mkdir(parents=True, exist_ok=True)
    incoming_path.write_bytes(data)

    version_number = next_storage_object_version_number(session, storage_object.id)
    version_id = new_id("storage_object_version")
    version_relative_path = storage_object_version_relative_path(
        storage_object,
        version_id=version_id,
        version_number=version_number,
    )
    version_path = (root / version_relative_path).resolve()
    if not version_path.is_relative_to(root):
        raise RuntimeError("Storage object version path escapes storage root.")

    previous_original_filename = storage_object.original_filename
    previous_content_type = storage_object.content_type
    previous_byte_size = storage_object.byte_size
    previous_sha256_hex = storage_object.sha256_hex
    previous_storage_path = storage_object.storage_path
    previous_file_updated_at = storage_object.file_updated_at

    version_path.parent.mkdir(parents=True, exist_ok=True)
    moved_current = False
    try:
        current_path.replace(version_path)
        moved_current = True
        new_path.parent.mkdir(parents=True, exist_ok=True)
        if new_path.is_file():
            new_path.unlink()
        incoming_path.replace(new_path)
    except Exception:
        if incoming_path.is_file():
            incoming_path.unlink()
        if moved_current and version_path.is_file() and not current_path.exists():
            version_path.replace(current_path)
        raise

    version = StorageObjectVersion(
        id=version_id,
        storage_object_id=storage_object.id,
        version_number=version_number,
        original_filename=previous_original_filename,
        content_type=previous_content_type,
        byte_size=previous_byte_size,
        sha256_hex=previous_sha256_hex,
        storage_path=version_relative_path.as_posix(),
        created_at=previous_file_updated_at,
    )
    session.add(version)

    storage_object.original_filename = filename
    storage_object.content_type = content_type
    storage_object.byte_size = len(data)
    storage_object.sha256_hex = new_sha256_hex
    storage_object.storage_path = new_relative_path.as_posix()
    storage_object.updated_at = now
    storage_object.file_updated_at = now
    storage_object.version += 1
    reassign_current_file_summaries_to_version(
        session,
        storage_object=storage_object,
        previous_version=version,
        previous_sha256_hex=previous_sha256_hex,
        now=now,
    )
    prepare_file_version_diff(
        session,
        storage_object=storage_object,
        previous_version=version,
        now=now,
    )
    details = {
        "previous_version_id": version.id,
        "previous_version_number": version.version_number,
        "previous_storage_path": previous_storage_path,
        "previous_file_updated_at": previous_file_updated_at,
        "new_filename": filename,
    }
    if operation_details is not None:
        details.update(operation_details)
    record_storage_operation(
        session,
        operation_type=operation_type,
        now=now,
        storage_object=storage_object,
        details=details,
    )
    return version


def delete_storage_object_file(
    storage_object: StorageObject,
    session: DatabaseSession,
) -> bool:
    object_path = storage_object_absolute_path(storage_object, session)
    if object_path.is_file():
        object_path.unlink()
        return True
    return False


def delete_storage_object_version_files(
    storage_object: StorageObject,
    session: DatabaseSession,
) -> int:
    root = storage_location_root(session, storage_object.location_id).resolve()
    versions = session.scalars(
        select(StorageObjectVersion).where(
            StorageObjectVersion.storage_object_id == storage_object.id
        )
    ).all()
    deleted_count = 0
    for version in versions:
        version_path = (root / version.storage_path).resolve()
        if not version_path.is_relative_to(root):
            raise RuntimeError("Storage object version path escapes storage root.")
        if version_path.is_file():
            version_path.unlink()
            deleted_count += 1
    return deleted_count


def delete_storage_object_versions_up_to(
    storage_object: StorageObject,
    selected_version: StorageObjectVersion,
    *,
    session: DatabaseSession,
    now: str,
) -> list[str]:
    root = storage_location_root(session, storage_object.location_id).resolve()
    versions = session.scalars(
        select(StorageObjectVersion)
        .where(StorageObjectVersion.storage_object_id == storage_object.id)
        .where(StorageObjectVersion.version_number <= selected_version.version_number)
        .order_by(StorageObjectVersion.version_number, StorageObjectVersion.id)
    ).all()
    deleted_version_ids: list[str] = []
    for version in versions:
        version_path = (root / version.storage_path).resolve()
        if not version_path.is_relative_to(root):
            raise RuntimeError("Storage object version path escapes storage root.")
        if version_path.is_file():
            version_path.unlink()
        deleted_version_ids.append(version.id)
        session.delete(version)
    storage_object.updated_at = now
    storage_object.version += 1
    record_storage_operation(
        session,
        operation_type="versions_deleted",
        now=now,
        storage_object=storage_object,
        details={
            "selected_version_id": selected_version.id,
            "selected_version_number": selected_version.version_number,
            "deleted_version_ids": deleted_version_ids,
            "deleted_version_count": len(deleted_version_ids),
        },
    )
    return deleted_version_ids


def delete_storage_object(
    storage_object: StorageObject,
    *,
    session: DatabaseSession,
    now: str,
) -> None:
    file_deleted = delete_storage_object_file(storage_object, session)
    version_files_deleted = delete_storage_object_version_files(storage_object, session)
    active_file_links = session.scalars(
        select(FileLink)
        .where(FileLink.storage_object_id == storage_object.id)
        .where(FileLink.status == "active")
    ).all()
    for file_link in active_file_links:
        file_link.status = "deleted"
        file_link.updated_at = now
        file_link.version += 1
    storage_object.status = "deleted"
    storage_object.updated_at = now
    storage_object.version += 1
    record_storage_operation(
        session,
        operation_type="deleted",
        now=now,
        storage_object=storage_object,
        details={
            "physical_file_deleted": file_deleted,
            "version_files_deleted": version_files_deleted,
            "deleted_file_link_count": len(active_file_links),
        },
    )


def delete_previous_contact_image(
    session: DatabaseSession,
    *,
    avatar_url: str | None,
    now: str,
) -> None:
    storage_object_id = storage_object_id_from_url(avatar_url)
    if storage_object_id is None:
        return
    storage_object = session.get(StorageObject, storage_object_id)
    if (
        storage_object is None
        or storage_object.status != "active"
        or storage_object.scope != "contact-images"
    ):
        return
    delete_storage_object(storage_object, session=session, now=now)


def save_file_icon_storage_object(
    session: DatabaseSession,
    *,
    filename: str | None,
    content_type: str,
    data_base64: str,
    now: str,
) -> StorageObject:
    resized_content_type, resized_data = prepare_file_icon_image(
        content_type=content_type,
        data_base64=data_base64,
    )
    return save_storage_object(
        session,
        scope="file-icons",
        filename=filename,
        content_type=resized_content_type,
        data=resized_data,
        now=now,
    )


def delete_file_icon_storage_object_if_unused(
    session: DatabaseSession,
    *,
    storage_object_id: str | None,
    setting_id: str | None = None,
    now: str,
) -> None:
    if storage_object_id is None:
        return
    statement = select(func.count(FileIconSetting.id)).where(
        FileIconSetting.storage_object_id == storage_object_id
    )
    if setting_id is not None:
        statement = statement.where(FileIconSetting.id != setting_id)
    if session.scalar(statement) != 0:
        return
    storage_object = session.get(StorageObject, storage_object_id)
    if (
        storage_object is None
        or storage_object.status != "active"
        or storage_object.scope != "file-icons"
    ):
        return
    delete_storage_object(storage_object, session=session, now=now)


def parse_paper_tags_json(value: str | None) -> list[str]:
    if value is None:
        return []
    try:
        parsed = json.loads(value)
    except json.JSONDecodeError:
        return []
    if not isinstance(parsed, list):
        return []
    return [item for item in parsed if isinstance(item, str)]


def normalize_paper_tags(values: list[str]) -> list[str]:
    tags: list[str] = []
    seen: set[str] = set()
    for raw_value in values:
        tag = raw_value.strip()
        normalized = tag.lower()
        if tag == "" or normalized in seen:
            continue
        seen.add(normalized)
        tags.append(tag)
    return tags


def paper_title_from_filename(filename: str | None) -> str:
    raw_name = (filename or "Untitled paper").strip() or "Untitled paper"
    return Path(raw_name).stem or raw_name


def normalize_bibtex_value(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip().strip(",")


def parse_ris_entry(ris: str) -> dict[str, object]:
    fields: dict[str, str] = {}
    authors: list[str] = []
    type_map = {
        "JOUR": "article",
        "CONF": "inproceedings",
        "CPAPER": "inproceedings",
        "BOOK": "book",
        "CHAP": "incollection",
        "THES": "phdthesis",
        "RPRT": "report",
    }
    entry_type = ""
    entry_key = ""
    for raw_line in ris.splitlines():
        line = raw_line.strip()
        if line == "" or line == "ER  -":
            continue
        match = re.match(r"(?P<tag>[A-Z0-9]{2})\s+-\s*(?P<value>.*)", line)
        if match is None:
            continue
        tag = match.group("tag")
        value = normalize_bibtex_value(match.group("value"))
        if value == "":
            continue
        if tag == "TY":
            entry_type = type_map.get(value.upper(), value.lower())
        elif tag in {"AU", "A1", "A2", "A3"}:
            authors.append(value)
        elif tag in {"TI", "T1"}:
            fields.setdefault("title", value)
        elif tag in {"JO", "JF", "JA", "T2"}:
            fields.setdefault("journal", value)
        elif tag in {"PY", "Y1"}:
            year_match = re.search(r"\d{4}", value)
            fields.setdefault("year", year_match.group(0) if year_match is not None else value)
        elif tag == "DO":
            fields["doi"] = value.removeprefix("https://doi.org/").removeprefix("http://doi.org/")
        elif tag in {"UR", "L1", "L2"}:
            fields.setdefault("url", value)
        elif tag in {"AB", "N2"}:
            fields.setdefault("abstract", value)
        elif tag == "ID":
            entry_key = value
        else:
            fields.setdefault(tag.lower(), value)
    if authors and "author" not in fields:
        fields["author"] = " and ".join(authors)
    if entry_key == "":
        entry_key = fields.get("doi") or fields.get("url") or fields.get("title", "")[:80]
    return {
        "entry_type": entry_type or "ris",
        "entry_key": entry_key,
        "fields": fields,
        "authors": authors,
    }



def is_nbib_text(raw_citation: str) -> bool:
    for raw_line in raw_citation.splitlines():
        if re.match(r"^(PMID|OWN|STAT|MHDA|CRDT|DP|JT|FAU|LID|AID)\s*-", raw_line.strip()):
            return True
    return False


def parse_nbib_entries(nbib: str) -> list[tuple[str, str]]:
    entries: list[tuple[str, str]] = []
    current_tag = ""
    current_value = ""
    for raw_line in nbib.splitlines():
        if raw_line.strip() == "":
            continue
        match = re.match(r"^(?P<tag>[A-Z0-9]{2,4})\s*-\s*(?P<value>.*)", raw_line)
        if match is not None:
            if current_tag != "":
                entries.append((current_tag, normalize_bibtex_value(current_value)))
            current_tag = match.group("tag")
            current_value = match.group("value")
            continue
        if current_tag != "" and raw_line.startswith(" "):
            current_value = f"{current_value} {raw_line.strip()}"
    if current_tag != "":
        entries.append((current_tag, normalize_bibtex_value(current_value)))
    return entries


def nbib_doi(value: str) -> str:
    cleaned = normalize_bibtex_value(value)
    match = re.search(r"(10\.\S+)", cleaned, flags=re.IGNORECASE)
    if match is None:
        return ""
    doi = match.group(1)
    doi = re.sub(r"\s*\[doi\]\s*$", "", doi, flags=re.IGNORECASE)
    return doi.rstrip(".,;")


def parse_nbib_entry(nbib: str) -> dict[str, object]:
    fields: dict[str, str] = {}
    short_authors: list[str] = []
    full_authors: list[str] = []
    for tag, value in parse_nbib_entries(nbib):
        if value == "":
            continue
        if tag == "PMID":
            fields["pmid"] = value
        elif tag == "TI":
            fields.setdefault("title", value)
        elif tag == "BTI":
            fields.setdefault("booktitle", value)
        elif tag == "AU":
            short_authors.append(value)
        elif tag == "FAU":
            full_authors.append(value)
        elif tag == "JT":
            fields.setdefault("journal", value)
        elif tag == "TA":
            fields.setdefault("journal", value)
        elif tag == "DP":
            year_match = re.search(r"\d{4}", value)
            fields["year"] = year_match.group(0) if year_match is not None else value
            fields["date"] = value
        elif tag == "VI":
            fields["volume"] = value
        elif tag == "IP":
            fields["number"] = value
        elif tag == "PG":
            fields["pages"] = value
        elif tag in {"LID", "AID"}:
            doi = nbib_doi(value)
            if doi != "":
                fields["doi"] = doi
        elif tag == "AB":
            fields.setdefault("abstract", value)
        elif tag == "PMC":
            fields["pmcid"] = value
        elif tag == "SO":
            fields.setdefault("source", value)
        else:
            fields.setdefault(tag.lower(), value)
    authors = full_authors if full_authors else short_authors
    if authors:
        fields["author"] = " and ".join(authors)
    if "url" not in fields and fields.get("pmid", "") != "":
        fields["url"] = f"https://pubmed.ncbi.nlm.nih.gov/{fields['pmid']}/"
    entry_key = fields.get("doi") or fields.get("pmid") or fields.get("title", "")[:80]
    return {
        "entry_type": "article",
        "entry_key": entry_key,
        "fields": fields,
        "authors": authors,
    }

def parse_citation_entry(raw_citation: str) -> dict[str, object]:
    if raw_citation.lstrip().startswith("@"):
        return parse_bibtex_entry(raw_citation)
    if is_nbib_text(raw_citation):
        return parse_nbib_entry(raw_citation)
    return parse_ris_entry(raw_citation)


def parse_bibtex_entry(bibtex: str) -> dict[str, object]:
    header_match = re.search(r"@(?P<type>[A-Za-z]+)\s*\{\s*(?P<key>[^,]+),", bibtex, re.DOTALL)
    entry_type = normalize_bibtex_value(header_match.group("type")) if header_match is not None else ""
    entry_key = normalize_bibtex_value(header_match.group("key")) if header_match is not None else ""
    start_index = header_match.end() if header_match is not None else 0
    end_index = bibtex.rfind("}")
    body = bibtex[start_index:end_index if end_index > start_index else len(bibtex)]
    fields: dict[str, str] = {}
    index = 0
    while index < len(body):
        while index < len(body) and body[index] in " \t\r\n,":
            index += 1
        name_start = index
        while index < len(body) and re.match(r"[A-Za-z0-9_:-]", body[index]):
            index += 1
        field_name = body[name_start:index].strip().lower()
        while index < len(body) and body[index].isspace():
            index += 1
        if field_name == "" or index >= len(body) or body[index] != "=":
            index += 1
            continue
        index += 1
        while index < len(body) and body[index].isspace():
            index += 1
        if index >= len(body):
            break
        if body[index] == "{":
            index += 1
            depth = 1
            value_start = index
            while index < len(body) and depth > 0:
                if body[index] == "{":
                    depth += 1
                elif body[index] == "}":
                    depth -= 1
                    if depth == 0:
                        break
                index += 1
            value = body[value_start:index]
            index += 1
        elif body[index] == '"':
            index += 1
            value_start = index
            escaped = False
            while index < len(body):
                character = body[index]
                if character == '"' and not escaped:
                    break
                escaped = character == "\\" and not escaped
                if character != "\\":
                    escaped = False
                index += 1
            value = body[value_start:index]
            index += 1
        else:
            value_start = index
            while index < len(body) and body[index] not in ",\n\r":
                index += 1
            value = body[value_start:index]
        fields[field_name] = normalize_bibtex_value(value)
    authors = [
        author.strip()
        for author in re.split(r"\s+and\s+", fields.get("author", ""), flags=re.IGNORECASE)
        if author.strip()
    ]
    return {
        "entry_type": entry_type,
        "entry_key": entry_key,
        "fields": fields,
        "authors": authors,
    }


def bibtex_field_value(bibtex: str, field_name: str) -> str | None:
    parsed = parse_citation_entry(bibtex)
    fields = parsed.get("fields", {})
    value = fields.get(field_name.lower()) if isinstance(fields, dict) else None
    return value if isinstance(value, str) and value != "" else None


async def read_paper_bibtex_upload_text(upload: UploadFile) -> str:
    if not is_citation_filename(upload.filename):
        raise json_error(422, "PAPER_CITATION_ONLY", "PaperShelf requires a .bib, .bibtex, .ris, or .nbib file.")
    data = await upload.read()
    await upload.close()
    if len(data) == 0:
        raise json_error(422, "VALIDATION_ERROR", "Citation file is empty.")
    if len(data) > MAX_PAPER_BIBTEX_BYTES:
        raise json_error(413, "PAYLOAD_TOO_LARGE", "Citation file is too large.")
    return data.decode("utf-8", errors="replace").strip()


def upsert_paper_bibtex_entry(
    session: DatabaseSession,
    paper: Paper,
    *,
    now: str | None = None,
) -> PaperBibtexEntry | None:
    raw_bibtex = paper.bibtex.strip()
    if raw_bibtex == "":
        return None
    timestamp = now or jst_iso()
    parsed = parse_citation_entry(raw_bibtex)
    fields = parsed.get("fields", {})
    if not isinstance(fields, dict):
        fields = {}
    authors = parsed.get("authors", [])
    if not isinstance(authors, list):
        authors = []
    entry = session.scalar(select(PaperBibtexEntry).where(PaperBibtexEntry.paper_id == paper.id).limit(1))
    if entry is None:
        entry = PaperBibtexEntry(
            id=new_id("paper_bibtex"),
            paper_id=paper.id,
            entry_type=str(parsed.get("entry_type", "")),
            entry_key=str(parsed.get("entry_key", "")),
            title=str(fields.get("title", "")),
            authors_json=json.dumps(authors, ensure_ascii=False),
            journal=str(fields.get("journal", "")),
            year=str(fields.get("year", "")),
            doi=str(fields.get("doi", "")),
            url=str(fields.get("url", "")),
            abstract=str(fields.get("abstract", "")),
            raw_bibtex=raw_bibtex,
            fields_json=json.dumps(fields, ensure_ascii=False),
            created_at=timestamp,
            updated_at=timestamp,
            version=1,
        )
        session.add(entry)
        return entry
    entry.entry_type = str(parsed.get("entry_type", ""))
    entry.entry_key = str(parsed.get("entry_key", ""))
    entry.title = str(fields.get("title", ""))
    entry.authors_json = json.dumps(authors, ensure_ascii=False)
    entry.journal = str(fields.get("journal", ""))
    entry.year = str(fields.get("year", ""))
    entry.doi = str(fields.get("doi", ""))
    entry.url = str(fields.get("url", ""))
    entry.abstract = str(fields.get("abstract", ""))
    entry.raw_bibtex = raw_bibtex
    entry.fields_json = json.dumps(fields, ensure_ascii=False)
    entry.updated_at = timestamp
    entry.version += 1
    return entry


def paper_bibtex_entry_data(entry: PaperBibtexEntry | None) -> dict[str, object] | None:
    if entry is None:
        return None
    try:
        authors = json.loads(entry.authors_json)
    except json.JSONDecodeError:
        authors = []
    try:
        fields = json.loads(entry.fields_json)
    except json.JSONDecodeError:
        fields = {}
    return {
        "id": entry.id,
        "paper_id": entry.paper_id,
        "entry_type": entry.entry_type,
        "entry_key": entry.entry_key,
        "title": entry.title,
        "authors": authors if isinstance(authors, list) else [],
        "journal": entry.journal,
        "year": entry.year,
        "doi": entry.doi,
        "url": entry.url,
        "abstract": entry.abstract,
        "raw_bibtex": entry.raw_bibtex,
        "fields": fields if isinstance(fields, dict) else {},
        "created_at": entry.created_at,
        "updated_at": entry.updated_at,
        "version": entry.version,
    }


def normalize_paper_journal_match(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def paper_journal_icon_data_url(*, content_type: str, data_base64: str) -> tuple[str, str]:
    resized_content_type, resized_data = prepare_file_icon_image(
        content_type=content_type,
        data_base64=data_base64,
    )
    encoded = base64.b64encode(resized_data).decode("ascii")
    return resized_content_type, f"data:{resized_content_type};base64,{encoded}"


def paper_journal_icon_setting_data(setting: PaperJournalIconSetting) -> dict[str, object]:
    return {
        "id": setting.id,
        "match_journal": setting.match_journal,
        "icon_filename": setting.icon_filename,
        "icon_content_type": setting.icon_content_type,
        "icon_url": setting.icon_data_url,
        "icon_data_url": setting.icon_data_url,
        "created_at": setting.created_at,
        "updated_at": setting.updated_at,
        "version": setting.version,
    }


def paper_journal_icon_for_name(
    session: DatabaseSession,
    journal_name: str,
) -> PaperJournalIconSetting | None:
    normalized_journal = journal_name.strip().lower()
    if normalized_journal == "":
        return None
    settings = session.scalars(
        select(PaperJournalIconSetting).order_by(
            PaperJournalIconSetting.created_at.asc(),
            PaperJournalIconSetting.id.asc(),
        )
    ).all()
    best: PaperJournalIconSetting | None = None
    best_length = -1
    for setting in settings:
        match_text = setting.match_journal.strip().lower()
        if match_text == "" or match_text not in normalized_journal:
            continue
        if len(match_text) > best_length:
            best = setting
            best_length = len(match_text)
    return best



def paper_data(paper: Paper, session: DatabaseSession) -> dict[str, object]:
    storage_object = session.get(StorageObject, paper.storage_object_id)
    bibtex_entry = upsert_paper_bibtex_entry(session, paper)
    journal_text = bibtex_entry.journal if bibtex_entry is not None else ""
    journal_icon = paper_journal_icon_for_name(session, journal_text)
    return {
        "id": paper.id,
        "storage_object_id": paper.storage_object_id,
        "storage_object": storage_object_data(storage_object, session) if storage_object is not None else None,
        "title": paper.title,
        "authors_text": paper.authors_text,
        "journal_text": journal_text,
        "journal_icon_url": journal_icon.icon_data_url if journal_icon is not None else None,
        "bibtex": bibtex_entry.raw_bibtex if bibtex_entry is not None else paper.bibtex,
        "summary": paper.summary,
        "bibtex_entry": paper_bibtex_entry_data(bibtex_entry),
        "tags": parse_paper_tags_json(paper.tags_json),
        "status": paper.status,
        "created_at": paper.created_at,
        "updated_at": paper.updated_at,
        "version": paper.version,
    }


def purge_bookshelf_material_object(
    storage_object: StorageObject,
    *,
    session: DatabaseSession,
) -> dict[str, object]:
    file_deleted = delete_storage_object_file(storage_object, session)
    version_files_deleted = delete_storage_object_version_files(storage_object, session)
    session.execute(delete(FileVersionDiff).where(FileVersionDiff.storage_object_id == storage_object.id))
    session.execute(delete(FileSummary).where(FileSummary.storage_object_id == storage_object.id))
    session.execute(delete(StorageObjectVersion).where(StorageObjectVersion.storage_object_id == storage_object.id))
    session.execute(delete(FileLink).where(FileLink.storage_object_id == storage_object.id))
    session.execute(
        delete(StorageOperationHistory).where(
            StorageOperationHistory.storage_object_id == storage_object.id
        )
    )
    purged_storage_object_id = storage_object.id
    session.delete(storage_object)
    return {
        "purged_storage_object_id": purged_storage_object_id,
        "physical_file_deleted": file_deleted,
        "version_files_deleted": version_files_deleted,
    }


def delete_managed_storage_object_for_request(
    storage_object: StorageObject,
    *,
    session: DatabaseSession,
    now: str,
) -> StorageObject | None:
    if storage_object.source_type == "mail_attachment":
        attachment = source_attachment_for_storage_object(session, storage_object)
        if attachment is None:
            raise json_error(
                409,
                "SOURCE_ATTACHMENT_NOT_FOUND",
                "Source mail attachment was not found.",
            )
        object_path = storage_object_absolute_path(storage_object, session)
        if not object_path.is_file():
            raise json_error(404, "NOT_FOUND", "Storage object file not found.")
        tmp_storage_object = save_storage_object(
            session,
            scope=GMAIL_ATTACHMENT_STORAGE_SCOPE,
            filename=storage_object.original_filename,
            content_type=storage_object.content_type or "application/octet-stream",
            data=object_path.read_bytes(),
            now=now,
            source_type="mail_attachment",
            source_message_id=storage_object.source_message_id,
        )
        attachment.storage_object_id = tmp_storage_object.id
        attachment.byte_size = tmp_storage_object.byte_size
        attachment.updated_at = now
        attachment.version += 1
        delete_storage_object(storage_object, session=session, now=now)
        return tmp_storage_object

    delete_storage_object(storage_object, session=session, now=now)
    return None


def record_storage_operation(
    session: DatabaseSession,
    *,
    operation_type: str,
    now: str,
    storage_object: StorageObject | None = None,
    actor: str = "system",
    details: dict[str, object] | None = None,
) -> None:
    session.add(
        StorageOperationHistory(
            id=new_id("storage_operation"),
            storage_object_id=storage_object.id if storage_object is not None else None,
            operation_type=operation_type,
            actor=actor,
            scope=storage_object.scope if storage_object is not None else None,
            original_filename=(
                storage_object.original_filename if storage_object is not None else None
            ),
            content_type=storage_object.content_type if storage_object is not None else None,
            byte_size=storage_object.byte_size if storage_object is not None else None,
            storage_path=storage_object.storage_path if storage_object is not None else None,
            source_type=storage_object.source_type if storage_object is not None else None,
            source_message_id=(
                storage_object.source_message_id if storage_object is not None else None
            ),
            directory_id=storage_object.directory_id if storage_object is not None else None,
            details_json=(
                json.dumps(details, ensure_ascii=True, sort_keys=True)
                if details is not None
                else None
            ),
            created_at=now,
        )
    )


def source_attachment_for_storage_object(
    session: DatabaseSession,
    storage_object: StorageObject,
) -> GmailMessageAttachment | None:
    if (
        storage_object.source_message_id is None
        or storage_object.original_filename is None
    ):
        return None
    attachments = session.scalars(
        select(GmailMessageAttachment)
        .where(GmailMessageAttachment.message_id == storage_object.source_message_id)
        .where(GmailMessageAttachment.filename == storage_object.original_filename)
        .order_by(GmailMessageAttachment.updated_at.desc(), GmailMessageAttachment.id.desc())
    ).all()
    for attachment in attachments:
        if attachment.storage_object_id == storage_object.id:
            return attachment
    for attachment in attachments:
        if attachment.byte_size in {0, storage_object.byte_size}:
            return attachment
    return attachments[0] if attachments else None


def decoded_zip_info_path(info: zipfile.ZipInfo) -> str:
    if info.flag_bits & 0x800:
        return info.filename
    try:
        return info.filename.encode("cp437").decode("cp932")
    except UnicodeError:
        return info.filename


def archive_tree_lines(
    directories: set[str],
    files: list[tuple[str, int]],
) -> list[str]:
    root: dict[str, object] = {"dirs": {}, "files": {}}
    for directory in directories:
        node = root
        for part in directory.strip("/").split("/"):
            if part == "":
                continue
            dirs = node["dirs"]
            assert isinstance(dirs, dict)
            node = dirs.setdefault(part, {"dirs": {}, "files": {}})
            assert isinstance(node, dict)
    for path, size in files:
        parts = [part for part in path.split("/") if part != ""]
        if not parts:
            continue
        node = root
        for part in parts[:-1]:
            dirs = node["dirs"]
            assert isinstance(dirs, dict)
            node = dirs.setdefault(part, {"dirs": {}, "files": {}})
            assert isinstance(node, dict)
        file_items = node["files"]
        assert isinstance(file_items, dict)
        file_items[parts[-1]] = size

    lines: list[str] = []

    def render(node: dict[str, object], prefix: str = "") -> None:
        dirs = node["dirs"]
        file_items = node["files"]
        assert isinstance(dirs, dict)
        assert isinstance(file_items, dict)
        entries = (
            [(name, child, True) for name, child in sorted(dirs.items())]
            + [(name, size, False) for name, size in sorted(file_items.items())]
        )
        for index, (name, value, is_directory) in enumerate(entries):
            if len(lines) >= MAX_ARCHIVE_TREE_ENTRIES:
                return
            is_last = index == len(entries) - 1
            connector = "└── " if is_last else "├── "
            if is_directory:
                lines.append(f"{prefix}{connector}{name}/")
                next_prefix = f"{prefix}{'    ' if is_last else '│   '}"
                assert isinstance(value, dict)
                render(value, next_prefix)
            else:
                assert isinstance(value, int)
                lines.append(f"{prefix}{connector}{name}  {format_bytes(value)}")

    render(root)
    return lines


def archive_tree_text(object_path: Path) -> dict[str, object]:
    try:
        with zipfile.ZipFile(object_path) as archive:
            infos = archive.infolist()
    except zipfile.BadZipFile as error:
        raise json_error(422, "INVALID_ARCHIVE", "Storage object is not a valid ZIP archive.") from error

    directories: set[str] = set()
    files: list[tuple[str, int]] = []
    total_uncompressed_size = 0
    for info in infos:
        path = decoded_zip_info_path(info).replace("\\", "/").strip("/")
        if path == "":
            continue
        parts = [part for part in path.split("/") if part not in {"", "."}]
        if not parts:
            continue
        for index in range(1, len(parts)):
            directories.add("/".join(parts[:index]) + "/")
        if info.is_dir():
            directories.add("/".join(parts) + "/")
            continue
        normalized_path = "/".join(parts)
        files.append((normalized_path, int(info.file_size or 0)))
        total_uncompressed_size += int(info.file_size or 0)

    entry_count = len(directories) + len(files)
    truncated = entry_count > MAX_ARCHIVE_TREE_ENTRIES
    lines = archive_tree_lines(directories, files)
    if truncated:
        lines.append(f"... truncated after {MAX_ARCHIVE_TREE_ENTRIES} entries")
    if not lines:
        lines.append("(empty archive)")
    return {
        "tree_text": "\n".join(lines),
        "entry_count": entry_count,
        "file_count": len(files),
        "directory_count": len(directories),
        "total_uncompressed_size": total_uncompressed_size,
        "truncated": truncated,
        "max_entries": MAX_ARCHIVE_TREE_ENTRIES,
    }


def message_header_text(message: EmailMessage, name: str) -> str | None:
    value = message.get(name)
    if value is None:
        return None
    text = str(value).strip()
    return text or None


def message_part_text(part: EmailMessage) -> str | None:
    try:
        content = part.get_content()
    except Exception:
        return None
    if isinstance(content, str):
        text = content.replace("\r\n", "\n").replace("\r", "\n").strip()
        return text or None
    return None


def first_header_address(header_value: str | None) -> str | None:
    if header_value is None or header_value.strip() == "":
        return None
    for _display_name, address in getaddresses([header_value]):
        if address.strip() == "":
            continue
        return address
    return None


def eml_preview_data(object_path: Path, *, session: DatabaseSession) -> dict[str, object]:
    if object_path.stat().st_size > MAX_EML_PREVIEW_BYTES:
        raise json_error(413, "PAYLOAD_TOO_LARGE", "EML file is too large to preview.")
    try:
        message = BytesParser(policy=policy.default).parsebytes(object_path.read_bytes())
    except Exception as error:
        raise json_error(422, "VALIDATION_ERROR", "Could not parse EML file.") from error

    body_plain: str | None = None
    body_html: str | None = None
    attachments: list[dict[str, object]] = []
    if message.is_multipart():
        preferred_plain = message.get_body(preferencelist=("plain",))
        preferred_html = message.get_body(preferencelist=("html",))
        if preferred_plain is not None:
            body_plain = message_part_text(preferred_plain)
        if preferred_html is not None:
            body_html = message_part_text(preferred_html)
        for part in message.walk():
            if part.is_multipart():
                continue
            disposition = (part.get_content_disposition() or "").lower()
            filename = part.get_filename()
            if disposition == "attachment" or filename:
                payload = part.get_payload(decode=True) or b""
                attachments.append(
                    {
                        "filename": filename or "(unnamed attachment)",
                        "content_type": part.get_content_type(),
                        "byte_size": len(payload),
                    }
                )
    elif message.get_content_type().lower() == "text/html":
        body_html = message_part_text(message)
    else:
        body_plain = message_part_text(message)

    if body_plain is None and body_html is None:
        fallback_body = message.get_body(preferencelist=("plain", "html"))
        if fallback_body is not None and fallback_body.get_content_type().lower() == "text/html":
            body_html = message_part_text(fallback_body)
        elif fallback_body is not None:
            body_plain = message_part_text(fallback_body)

    from_header = message_header_text(message, "from")
    reply_to_header = message_header_text(message, "reply-to")
    from_address = first_header_address(from_header)
    reply_to_address = first_header_address(reply_to_header)
    from caseclosed.mails import contact_summary
    from caseclosed.mails import display_sender_contact_for_addresses

    sender_contact = (
        display_sender_contact_for_addresses(
            session,
            from_address=from_address,
            reply_to_address=reply_to_address,
        )
        if from_address is not None
        else None
    )
    return {
        "subject": message_header_text(message, "subject"),
        "from": from_header,
        "to": message_header_text(message, "to"),
        "cc": message_header_text(message, "cc"),
        "date": message_header_text(message, "date"),
        "reply_to": reply_to_header,
        "message_id": message_header_text(message, "message-id"),
        "body_text": body_plain,
        "body_html": body_html,
        "attachments": attachments,
        "sender_contact": (
            contact_summary(sender_contact, session=session) if sender_contact is not None else None
        ),
    }


def is_textual_storage_file(content_type: str | None, filename: str | None) -> bool:
    normalized_content_type = (content_type or "").split(";", 1)[0].strip().lower()
    if normalized_content_type.startswith("text/"):
        return True
    if normalized_content_type in {
        "application/json",
        "application/xml",
        "application/x-yaml",
        "application/yaml",
        "application/csv",
        "application/javascript",
        "application/x-tex",
    }:
        return True
    suffix = Path(filename or "").suffix.lower()
    return suffix in {
        ".txt",
        ".md",
        ".csv",
        ".tsv",
        ".json",
        ".xml",
        ".yaml",
        ".yml",
        ".py",
        ".js",
        ".ts",
        ".tsx",
        ".html",
        ".css",
        ".tex",
        ".bib",
        ".log",
    }


def normalized_file_extension(filename: str | None) -> str:
    return Path(filename or "").suffix.lower()


def is_pdf_storage_file(content_type: str | None, filename: str | None) -> bool:
    normalized_content_type = (content_type or "").split(";", 1)[0].strip().lower()
    return normalized_content_type == "application/pdf" or normalized_file_extension(filename) == ".pdf"


def is_docx_storage_file(content_type: str | None, filename: str | None) -> bool:
    normalized_content_type = (content_type or "").split(";", 1)[0].strip().lower()
    return (
        normalized_content_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or normalized_file_extension(filename) == ".docx"
    )


def is_xlsx_storage_file(content_type: str | None, filename: str | None) -> bool:
    normalized_content_type = (content_type or "").split(";", 1)[0].strip().lower()
    return (
        normalized_content_type
        == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        or normalized_file_extension(filename) in {".xlsx", ".xmlx"}
    )


def is_eml_storage_file(content_type: str | None, filename: str | None) -> bool:
    normalized_content_type = (content_type or "").split(";", 1)[0].strip().lower()
    return normalized_content_type == "message/rfc822" or normalized_file_extension(filename) == ".eml"


def best_effort_decode_text(data: bytes) -> str:
    best_text = ""
    best_score: int | None = None
    for encoding in ("utf-8-sig", "utf-8", "cp932", "shift_jis", "euc_jp"):
        try:
            text_value = data.decode(encoding)
        except UnicodeDecodeError:
            continue
        score = text_value.count("\ufffd") * 100
        score += sum(
            1
            for char in text_value
            if ord(char) < 32 and char not in {"\n", "\r", "\t"}
        )
        if best_score is None or score < best_score:
            best_text = text_value
            best_score = score
    if best_score is None:
        return data.decode("utf-8", errors="replace")
    return best_text


def limited_text(text_value: str, limitations: list[str]) -> tuple[str, bool]:
    if len(text_value) <= MAX_FILE_SUMMARY_INPUT_CHARS:
        return text_value, False
    limitations.append("Extracted text exceeded the character limit.")
    return text_value[:MAX_FILE_SUMMARY_INPUT_CHARS], True


def extract_plain_text_for_summary(path: Path) -> FileTextExtraction:
    raw_data = path.read_bytes()
    limitations: list[str] = []
    truncated = len(raw_data) > MAX_FILE_SUMMARY_INPUT_BYTES
    if truncated:
        raw_data = raw_data[:MAX_FILE_SUMMARY_INPUT_BYTES]
        limitations.append("File text was truncated before LLM digest generation.")
    source_text = best_effort_decode_text(raw_data)
    source_text, character_truncated = limited_text(source_text, limitations)
    return FileTextExtraction(
        source_kind="text",
        read_scope="text_content",
        source_text=source_text,
        truncated=truncated or character_truncated,
        limitations=limitations,
    )


def extract_archive_tree_for_summary(path: Path) -> FileTextExtraction:
    archive_info = archive_tree_text(path)
    limitations: list[str] = []
    truncated = bool(archive_info["truncated"])
    if truncated:
        limitations.append("Archive tree was truncated.")
    return FileTextExtraction(
        source_kind="archive_tree",
        read_scope="zip_entry_tree",
        source_text=str(archive_info["tree_text"]),
        truncated=truncated,
        limitations=limitations,
    )


def pdf_literal_text(value: str) -> str:
    value = value.replace("\\(", "(").replace("\\)", ")").replace("\\\\", "\\")
    value = re.sub(
        r"\\([0-7]{1,3})",
        lambda match: chr(int(match.group(1), 8)),
        value,
    )
    return value


def primitive_pdf_text(path: Path) -> str:
    data = path.read_bytes()
    chunks: list[str] = []
    for match in re.finditer(rb"stream\r?\n(.*?)\r?\nendstream", data, flags=re.S):
        stream_data = match.group(1).strip(b"\r\n")
        dictionary_prefix = data[max(0, match.start() - 600) : match.start()]
        if b"/FlateDecode" in dictionary_prefix:
            try:
                stream_data = zlib.decompress(stream_data)
            except zlib.error:
                continue
        text = stream_data.decode("latin-1", errors="ignore")
        for literal in re.findall(r"\((?:\\.|[^\\)])*\)\s*Tj", text):
            chunks.append(pdf_literal_text(literal.rsplit(")", 1)[0][1:]))
        for array_text in re.findall(r"\[(.*?)\]\s*TJ", text, flags=re.S):
            chunks.extend(
                pdf_literal_text(item[1:-1])
                for item in re.findall(r"\((?:\\.|[^\\)])*\)", array_text)
            )
    return "\n".join(chunk for chunk in chunks if chunk.strip())


def extract_pdf_with_fitz(path: Path) -> tuple[str, int] | None:
    try:
        import fitz  # type: ignore[import-not-found]
    except ImportError:
        return None
    document = fitz.open(path)
    try:
        return "\n".join(page.get_text("text") for page in document), len(document)
    finally:
        document.close()


def extract_pdf_with_pypdf(path: Path) -> tuple[str, int] | None:
    try:
        from pypdf import PdfReader  # type: ignore[import-not-found]
    except ImportError:
        try:
            from PyPDF2 import PdfReader  # type: ignore[import-not-found,no-redef]
        except ImportError:
            return None
    reader = PdfReader(str(path))
    texts: list[str] = []
    for page in reader.pages:
        try:
            texts.append(page.extract_text() or "")
        except Exception:
            texts.append("")
    return "\n".join(texts), len(reader.pages)


def extract_pdf_for_summary(path: Path) -> FileTextExtraction:
    limitations: list[str] = []
    page_count: int | None = None
    extracted = extract_pdf_with_fitz(path)
    extractor = "fitz"
    if extracted is None:
        extracted = extract_pdf_with_pypdf(path)
        extractor = "pypdf"
    if extracted is None:
        source_text = primitive_pdf_text(path)
        extractor = "primitive_pdf"
    else:
        source_text, page_count = extracted

    source_text = source_text.strip()
    if source_text == "":
        limitations.append("PDF text extraction returned no text; OCR may be required.")
    if extractor == "primitive_pdf":
        limitations.append("Used primitive PDF text extraction because PDF libraries are unavailable.")
    source_text, truncated = limited_text(source_text, limitations)
    if page_count is not None:
        limitations.append(f"PDF page count: {page_count}.")
    return FileTextExtraction(
        source_kind="pdf_text" if source_text else "pdf_no_text",
        read_scope=f"pdf_text_extraction:{extractor}",
        source_text=source_text,
        truncated=truncated,
        limitations=limitations,
    )


def pdf_text_quality_data(path: Path) -> dict[str, object]:
    limitations: list[str] = []
    page_count: int | None = None
    extracted = extract_pdf_with_fitz(path)
    extractor = "fitz"
    if extracted is None:
        extracted = extract_pdf_with_pypdf(path)
        extractor = "pypdf"
    if extracted is None:
        source_text = primitive_pdf_text(path)
        extractor = "primitive_pdf"
    else:
        source_text, page_count = extracted
    text_value = source_text.strip()
    compact_text = re.sub(r"\s+", "", text_value)
    native_char_count = len(compact_text)
    replacement_ratio = (text_value.count("�") / max(len(text_value), 1)) if text_value else 0.0
    useful_chars = re.findall(r"[A-Za-z0-9ぁ-んァ-ン一-龯々〆〤ー]", text_value)
    useful_ratio = len(useful_chars) / max(native_char_count, 1)
    whitespace_ratio = (sum(1 for char in text_value if char.isspace()) / max(len(text_value), 1)) if text_value else 1.0
    minimum_chars = max(40, (page_count or 1) * 24)
    reasons: list[str] = []
    if native_char_count == 0:
        reasons.append("no_native_text")
    elif native_char_count < minimum_chars:
        reasons.append("too_little_native_text")
    if replacement_ratio > 0.01:
        reasons.append("replacement_character_noise")
    if native_char_count >= 40 and useful_ratio < 0.45:
        reasons.append("low_useful_character_ratio")
    if len(text_value) >= 100 and whitespace_ratio > 0.58:
        reasons.append("excessive_whitespace")
    if extractor == "primitive_pdf":
        reasons.append("primitive_extractor")
    text_quality = "missing" if native_char_count == 0 else "suspect" if reasons else "good"
    status = "native_text" if text_quality == "good" else "ocr_needed"
    return {
        "status": status,
        "text_quality": text_quality,
        "page_count": page_count,
        "native_char_count": native_char_count,
        "read_scope": f"pdf_text_extraction:{extractor}",
        "reasons": reasons,
        "replacement_ratio": replacement_ratio,
        "useful_ratio": useful_ratio,
        "whitespace_ratio": whitespace_ratio,
        "limitations": limitations,
    }


def pdf_ocr_status_data(status: StoragePdfOcrStatus) -> dict[str, object]:
    try:
        details = json.loads(status.details_json or "{}")
    except json.JSONDecodeError:
        details = {}
    return {
        "id": status.id,
        "storage_object_id": status.storage_object_id,
        "status": status.status,
        "text_quality": status.text_quality,
        "page_count": status.page_count,
        "native_char_count": status.native_char_count,
        "ocr_engine": status.ocr_engine,
        "error_message": status.error_message,
        "details": details if isinstance(details, dict) else {},
        "created_at": status.created_at,
        "updated_at": status.updated_at,
        "version": status.version,
    }


def upsert_pdf_ocr_status(
    session: DatabaseSession,
    storage_object: StorageObject,
    *,
    now: str,
    status: str,
    text_quality: str,
    page_count: int | None,
    native_char_count: int,
    ocr_engine: str | None = None,
    error_message: str | None = None,
    details: dict[str, object] | None = None,
) -> StoragePdfOcrStatus:
    row = session.scalar(
        select(StoragePdfOcrStatus)
        .where(StoragePdfOcrStatus.storage_object_id == storage_object.id)
        .limit(1)
    )
    details_json = json.dumps(details or {}, ensure_ascii=False, sort_keys=True)
    if row is None:
        row = StoragePdfOcrStatus(
            id=new_id("storage_pdf_ocr"),
            storage_object_id=storage_object.id,
            status=status,
            text_quality=text_quality,
            page_count=page_count,
            native_char_count=native_char_count,
            ocr_engine=ocr_engine,
            error_message=error_message,
            details_json=details_json,
            created_at=now,
            updated_at=now,
            version=1,
        )
        session.add(row)
        return row
    row.status = status
    row.text_quality = text_quality
    row.page_count = page_count
    row.native_char_count = native_char_count
    row.ocr_engine = ocr_engine
    row.error_message = error_message
    row.details_json = details_json
    row.updated_at = now
    row.version += 1
    return row


def ensure_pdf_ocr_status(
    session: DatabaseSession,
    storage_object: StorageObject,
    *,
    refresh: bool = False,
) -> StoragePdfOcrStatus:
    existing = session.scalar(
        select(StoragePdfOcrStatus)
        .where(StoragePdfOcrStatus.storage_object_id == storage_object.id)
        .limit(1)
    )
    if existing is not None and not refresh:
        try:
            details = json.loads(existing.details_json or "{}")
        except json.JSONDecodeError:
            details = {}
        if details.get("sha256_hex") == storage_object.sha256_hex:
            return existing
    now = jst_iso()
    object_path = storage_object_absolute_path(storage_object, session)
    quality = pdf_text_quality_data(object_path)
    details = {
        "sha256_hex": storage_object.sha256_hex,
        "read_scope": quality["read_scope"],
        "reasons": quality["reasons"],
        "replacement_ratio": quality["replacement_ratio"],
        "useful_ratio": quality["useful_ratio"],
        "whitespace_ratio": quality["whitespace_ratio"],
        "limitations": quality["limitations"],
    }
    return upsert_pdf_ocr_status(
        session,
        storage_object,
        now=now,
        status=str(quality["status"]),
        text_quality=str(quality["text_quality"]),
        page_count=quality["page_count"] if isinstance(quality["page_count"], int) else None,
        native_char_count=int(quality["native_char_count"]),
        details=details,
    )


def require_pdf_storage_object_for_ocr(
    session: DatabaseSession,
    storage_object_id: str,
) -> StorageObject:
    storage_object = session.get(StorageObject, storage_object_id)
    if (
        storage_object is None
        or storage_object.status != "active"
        or storage_object.scope != "managed"
        or not is_pdf_storage_file(storage_object.content_type, storage_object.original_filename)
    ):
        raise json_error(404, "NOT_FOUND", "PDF storage object not found.")
    return storage_object


def run_ocrmypdf_for_storage_object(
    session: DatabaseSession,
    storage_object: StorageObject,
    *,
    now: str,
) -> tuple[StorageObjectVersion | None, StoragePdfOcrStatus]:
    ocrmypdf_path = shutil.which("ocrmypdf")
    if ocrmypdf_path is None:
        status_row = upsert_pdf_ocr_status(
            session,
            storage_object,
            now=now,
            status="ocr_failed",
            text_quality="unknown",
            page_count=None,
            native_char_count=0,
            error_message="ocrmypdf is not installed.",
            details={"sha256_hex": storage_object.sha256_hex, "reason": "ocrmypdf_unavailable"},
        )
        session.commit()
        raise json_error(503, "OCR_UNAVAILABLE", "ocrmypdf is not installed.")
    object_path = storage_object_absolute_path(storage_object, session)
    with tempfile.TemporaryDirectory(prefix="caseclosed-ocr-") as tmp_dir:
        tmp_path = Path(tmp_dir)
        output_path = tmp_path / "ocr.pdf"
        command = [
            ocrmypdf_path,
            "--force-ocr",
            "--deskew",
            "--optimize",
            "0",
            "-l",
            "jpn+eng",
            str(object_path),
            str(output_path),
        ]
        try:
            result = subprocess.run(command, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=600)
        except subprocess.TimeoutExpired as error:
            upsert_pdf_ocr_status(
                session,
                storage_object,
                now=now,
                status="ocr_failed",
                text_quality="unknown",
                page_count=None,
                native_char_count=0,
                ocr_engine="ocrmypdf",
                error_message="OCR timed out.",
                details={"sha256_hex": storage_object.sha256_hex, "reason": "timeout"},
            )
            session.commit()
            raise json_error(504, "OCR_TIMEOUT", "OCR timed out.") from error
        except subprocess.CalledProcessError as error:
            detail = (error.stderr or error.stdout or b"").decode("utf-8", errors="replace").strip()
            message = detail[:1200] or "OCR failed."
            upsert_pdf_ocr_status(
                session,
                storage_object,
                now=now,
                status="ocr_failed",
                text_quality="unknown",
                page_count=None,
                native_char_count=0,
                ocr_engine="ocrmypdf",
                error_message=message,
                details={"sha256_hex": storage_object.sha256_hex, "reason": "process_failed"},
            )
            session.commit()
            raise json_error(500, "OCR_FAILED", message) from error
        if not output_path.is_file() or output_path.stat().st_size == 0:
            upsert_pdf_ocr_status(
                session,
                storage_object,
                now=now,
                status="ocr_failed",
                text_quality="unknown",
                page_count=None,
                native_char_count=0,
                ocr_engine="ocrmypdf",
                error_message="OCR did not produce a PDF.",
                details={"sha256_hex": storage_object.sha256_hex, "reason": "empty_output"},
            )
            session.commit()
            raise json_error(500, "OCR_FAILED", "OCR did not produce a PDF.")
        output_data = output_path.read_bytes()
    version = update_storage_object_from_bytes(
        session,
        storage_object,
        data=output_data,
        filename=storage_object.original_filename,
        content_type="application/pdf",
        now=now,
        operation_type="pdf_ocr_completed",
        operation_details={"engine": "ocrmypdf"},
    )
    quality = pdf_text_quality_data(storage_object_absolute_path(storage_object, session))
    status_row = upsert_pdf_ocr_status(
        session,
        storage_object,
        now=now,
        status="ocr_done",
        text_quality=str(quality["text_quality"]),
        page_count=quality["page_count"] if isinstance(quality["page_count"], int) else None,
        native_char_count=int(quality["native_char_count"]),
        ocr_engine="ocrmypdf",
        error_message=None,
        details={
            "sha256_hex": storage_object.sha256_hex,
            "read_scope": quality["read_scope"],
            "reasons": quality["reasons"],
            "ocr_stdout": result.stdout.decode("utf-8", errors="replace")[:1200],
            "ocr_stderr": result.stderr.decode("utf-8", errors="replace")[:1200],
        },
    )
    return version, status_row


def extract_docx_for_summary(path: Path) -> FileTextExtraction:
    try:
        from docx import Document  # type: ignore[import-not-found]
    except ImportError:
        return FileTextExtraction(
            source_kind="docx_no_text",
            read_scope="docx_text_extraction:unavailable",
            source_text="",
            truncated=False,
            limitations=["DOCX text extraction requires python-docx."],
        )
    document = Document(str(path))
    parts: list[str] = []
    paragraph_text = "\n".join(
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip() != ""
    )
    if paragraph_text != "":
        parts.append(paragraph_text)
    for table_index, table in enumerate(document.tables, start=1):
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.replace("\r\n", "\n").replace("\r", "\n").strip() for cell in row.cells]
            if any(cell != "" for cell in cells):
                rows.append(" | ".join(cells))
        if rows:
            parts.append(f"[Table {table_index}]\n" + "\n".join(rows))
    text = "\n\n".join(parts)
    limitations: list[str] = []
    text, truncated = limited_text(text, limitations)
    return FileTextExtraction(
        source_kind="docx_text" if text.strip() else "docx_no_text",
        read_scope="docx_paragraph_and_table_text",
        source_text=text,
        truncated=truncated,
        limitations=limitations,
    )


def xml_text(element: ET.Element) -> str:
    return "".join(element.itertext()).strip()


def xlsx_shared_strings(archive: zipfile.ZipFile) -> list[str]:
    try:
        raw_data = archive.read("xl/sharedStrings.xml")
    except KeyError:
        return []
    try:
        root = ET.fromstring(raw_data)
    except ET.ParseError:
        return []
    strings: list[str] = []
    for item in root.iter():
        if item.tag.endswith("}si") or item.tag == "si":
            strings.append(xml_text(item))
    return strings


def xlsx_cell_text(cell: ET.Element, shared_strings: list[str]) -> str:
    cell_type = cell.attrib.get("t")
    if cell_type == "inlineStr":
        for child in cell:
            if child.tag.endswith("}is") or child.tag == "is":
                return xml_text(child)
        return xml_text(cell)
    value_element: ET.Element | None = None
    for child in cell:
        if child.tag.endswith("}v") or child.tag == "v":
            value_element = child
            break
    value = (value_element.text or "").strip() if value_element is not None else ""
    if cell_type == "s":
        try:
            return shared_strings[int(value)]
        except (ValueError, IndexError):
            return value
    return value


def extract_xlsx_for_summary(path: Path) -> FileTextExtraction:
    limitations: list[str] = []
    try:
        archive = zipfile.ZipFile(path)
    except zipfile.BadZipFile:
        return FileTextExtraction(
            source_kind="xlsx_no_text",
            read_scope="xlsx_csv_extraction:invalid_zip",
            source_text="",
            truncated=False,
            limitations=["XLSX file is not a valid ZIP package."],
        )
    with archive:
        shared_strings = xlsx_shared_strings(archive)
        sheet_names = sorted(
            name
            for name in archive.namelist()
            if re.fullmatch(r"xl/worksheets/sheet\d+\.xml", name)
        )
        parts: list[str] = []
        for sheet_index, sheet_name in enumerate(sheet_names, start=1):
            try:
                root = ET.fromstring(archive.read(sheet_name))
            except (KeyError, ET.ParseError):
                limitations.append(f"Could not parse {sheet_name}.")
                continue
            output = StringIO()
            writer = csv.writer(output, lineterminator="\n")
            for row in root.iter():
                if not (row.tag.endswith("}row") or row.tag == "row"):
                    continue
                values = [
                    xlsx_cell_text(cell, shared_strings)
                    for cell in row
                    if cell.tag.endswith("}c") or cell.tag == "c"
                ]
                if any(value != "" for value in values):
                    writer.writerow(values)
            csv_text = output.getvalue().strip()
            if csv_text != "":
                parts.append(f"[Sheet {sheet_index}: {sheet_name}]\n{csv_text}")
        text = "\n\n".join(parts)
    text, truncated = limited_text(text, limitations)
    if text.strip() == "":
        limitations.append("XLSX text extraction returned no cell values.")
    return FileTextExtraction(
        source_kind="xlsx_csv" if text.strip() else "xlsx_no_text",
        read_scope="xlsx_sheet_xml_to_csv",
        source_text=text,
        truncated=truncated,
        limitations=limitations,
    )


def plain_text_from_html(value: str | None) -> str | None:
    if value is None:
        return None
    text = re.sub(r"(?is)<(script|style).*?>.*?</\1>", "", value)
    text = re.sub(r"(?i)<br\s*/?>", "\n", text)
    text = re.sub(r"(?i)</p\s*>", "\n", text)
    text = re.sub(r"<[^>]+>", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    stripped = text.strip()
    return stripped or None


def extract_eml_for_summary(path: Path, *, session: DatabaseSession) -> FileTextExtraction:
    limitations: list[str] = []
    try:
        preview = eml_preview_data(path, session=session)
    except Exception as error:
        return FileTextExtraction(
            source_kind="eml_no_text",
            read_scope="eml_parse_failed",
            source_text="",
            truncated=False,
            limitations=[f"EML parsing failed: {error}"],
        )
    body_text = preview.get("body_text")
    if not isinstance(body_text, str) or body_text.strip() == "":
        body_html = preview.get("body_html")
        body_text = plain_text_from_html(body_html if isinstance(body_html, str) else None) or ""
        if body_text != "":
            limitations.append("Plain text body was unavailable; converted HTML body to text.")
    attachments_payload = preview.get("attachments")
    attachments = attachments_payload if isinstance(attachments_payload, list) else []
    lines = [
        f"Subject: {preview.get('subject') or ''}",
        f"From: {preview.get('from') or ''}",
        f"Reply-To: {preview.get('reply_to') or ''}",
        f"To: {preview.get('to') or ''}",
        f"Cc: {preview.get('cc') or ''}",
        f"Date: {preview.get('date') or ''}",
        f"Message-ID: {preview.get('message_id') or ''}",
        "",
        "Body:",
        str(body_text or ""),
    ]
    if attachments:
        lines.extend(["", "Attachments:"])
        for attachment in attachments:
            if not isinstance(attachment, dict):
                continue
            lines.append(
                "- "
                f"{attachment.get('filename') or ''} "
                f"({attachment.get('content_type') or 'unknown'}, "
                f"{attachment.get('byte_size') or 0} bytes)"
            )
    text, truncated = limited_text("\n".join(lines), limitations)
    return FileTextExtraction(
        source_kind="eml_text" if text.strip() else "eml_no_text",
        read_scope="eml_headers_body_attachments",
        source_text=text,
        truncated=truncated,
        limitations=limitations,
    )


def extract_file_text_for_summary(
    path: Path,
    *,
    content_type: str | None,
    filename: str | None,
    session: DatabaseSession | None = None,
) -> FileTextExtraction:
    if is_pdf_storage_file(content_type, filename):
        return extract_pdf_for_summary(path)
    if is_docx_storage_file(content_type, filename):
        return extract_docx_for_summary(path)
    if is_xlsx_storage_file(content_type, filename):
        return extract_xlsx_for_summary(path)
    if is_eml_storage_file(content_type, filename):
        if session is None:
            return FileTextExtraction(
                source_kind="eml_no_text",
                read_scope="eml_parse:session_unavailable",
                source_text="",
                truncated=False,
                limitations=["EML text extraction requires a database session."],
            )
        return extract_eml_for_summary(path, session=session)
    if zipfile.is_zipfile(path):
        return extract_archive_tree_for_summary(path)
    if is_textual_storage_file(content_type, filename):
        return extract_plain_text_for_summary(path)
    return FileTextExtraction(
        source_kind="metadata",
        read_scope="metadata_only",
        source_text="",
        truncated=False,
        limitations=["File body extraction is not supported for this file type yet."],
    )


def storage_summary_source(
    storage_object: StorageObject,
    version: StorageObjectVersion | None,
    session: DatabaseSession,
) -> dict[str, object]:
    if version is None:
        path = storage_object_absolute_path(storage_object, session)
        filename = storage_object.original_filename
        content_type = storage_object.content_type
        byte_size = storage_object.byte_size
        sha256_hex = storage_object.sha256_hex
        version_id = None
    else:
        path = storage_object_version_absolute_path(storage_object, version, session)
        filename = version.original_filename
        content_type = version.content_type
        byte_size = version.byte_size
        sha256_hex = version.sha256_hex
        version_id = version.id

    if not path.is_file():
        raise json_error(404, "NOT_FOUND", "Storage object file not found.")
    extraction = extract_file_text_for_summary(
        path,
        content_type=content_type,
        filename=filename,
        session=session,
    )

    return {
        "storage_object_id": storage_object.id,
        "storage_object_version_id": version_id,
        "filename": filename,
        "content_type": content_type,
        "byte_size": byte_size,
        "sha256_hex": sha256_hex,
        "source_text": extraction.source_text,
        "source_kind": extraction.source_kind,
        "read_scope": extraction.read_scope,
        "truncated": extraction.truncated,
        "limitations": extraction.limitations,
    }


def latest_matching_file_summary(
    session: DatabaseSession,
    *,
    storage_object_id: str,
    storage_object_version_id: str | None,
    source_sha256_hex: str,
) -> FileSummary | None:
    query = select(FileSummary).where(
        FileSummary.storage_object_id == storage_object_id,
        FileSummary.source_sha256_hex == source_sha256_hex,
        FileSummary.summary_type == "llm_digest",
    )
    if storage_object_version_id is None:
        query = query.where(FileSummary.storage_object_version_id.is_(None))
    else:
        query = query.where(
            FileSummary.storage_object_version_id == storage_object_version_id
        )
    return session.scalar(
        query.order_by(FileSummary.updated_at.desc(), FileSummary.created_at.desc()).limit(1)
    )


def latest_any_file_summary(
    session: DatabaseSession,
    *,
    storage_object_id: str,
) -> FileSummary | None:
    return session.scalar(
        select(FileSummary)
        .where(FileSummary.storage_object_id == storage_object_id)
        .where(FileSummary.summary_type == "llm_digest")
        .order_by(FileSummary.updated_at.desc(), FileSummary.created_at.desc())
        .limit(1)
    )


def display_file_summary(
    session: DatabaseSession,
    *,
    storage_object_id: str,
    storage_object_version_id: str | None,
    source_sha256_hex: str,
) -> tuple[FileSummary | None, bool, str | None]:
    exact_summary = latest_matching_file_summary(
        session,
        storage_object_id=storage_object_id,
        storage_object_version_id=storage_object_version_id,
        source_sha256_hex=source_sha256_hex,
    )
    if exact_summary is not None:
        return exact_summary, False, None
    if storage_object_version_id is not None:
        return None, False, None
    stale_summary = latest_any_file_summary(
        session,
        storage_object_id=storage_object_id,
    )
    if stale_summary is None:
        return None, False, None
    return stale_summary, True, "Digest was prepared from an older file version."


def reassign_current_file_summaries_to_version(
    session: DatabaseSession,
    *,
    storage_object: StorageObject,
    previous_version: StorageObjectVersion,
    previous_sha256_hex: str,
    now: str,
) -> None:
    summaries = session.scalars(
        select(FileSummary)
        .where(FileSummary.storage_object_id == storage_object.id)
        .where(FileSummary.storage_object_version_id.is_(None))
        .where(FileSummary.source_sha256_hex == previous_sha256_hex)
        .where(FileSummary.summary_type == "llm_digest")
    ).all()
    for summary in summaries:
        summary.storage_object_version_id = previous_version.id
        summary.updated_at = now
        summary.version += 1


def compact_diff_lines(lines: list[str], *, limit: int = 12) -> list[str]:
    compacted = []
    for line in lines:
        normalized = re.sub(r"\s+", " ", line).strip()
        if normalized == "":
            continue
        compacted.append(normalized[:240])
        if len(compacted) >= limit:
            break
    return compacted


def prepare_file_version_diff(
    session: DatabaseSession,
    *,
    storage_object: StorageObject,
    previous_version: StorageObjectVersion,
    now: str,
) -> FileVersionDiff | None:
    try:
        previous_source = storage_summary_source(storage_object, previous_version, session)
        current_source = storage_summary_source(storage_object, None, session)
    except Exception:
        return None

    previous_text = str(previous_source.get("source_text") or "")
    current_text = str(current_source.get("source_text") or "")
    limitations = []
    limitations.extend(previous_source.get("limitations") or [])
    limitations.extend(current_source.get("limitations") or [])

    if previous_text.strip() == "" or current_text.strip() == "":
        diff_kind = "unavailable"
        added_lines: list[str] = []
        removed_lines: list[str] = []
        summary_text = "Text diff was not available for one or both versions."
    else:
        previous_lines = previous_text.splitlines()
        current_lines = current_text.splitlines()
        diff = list(difflib.ndiff(previous_lines, current_lines))
        removed_lines = compact_diff_lines([line[2:] for line in diff if line.startswith("- ")])
        added_lines = compact_diff_lines([line[2:] for line in diff if line.startswith("+ ")])
        if not added_lines and not removed_lines:
            diff_kind = "unchanged_text"
            summary_text = "Extracted text is unchanged."
        else:
            diff_kind = "line_diff"
            summary_text = (
                f"Added {len(added_lines)} notable line(s), "
                f"removed {len(removed_lines)} notable line(s)."
            )

    current_sha256_hex = str(current_source["sha256_hex"])
    existing = session.scalar(
        select(FileVersionDiff)
        .where(FileVersionDiff.storage_object_id == storage_object.id)
        .where(FileVersionDiff.previous_version_id == previous_version.id)
        .where(FileVersionDiff.current_sha256_hex == current_sha256_hex)
        .limit(1)
    )
    coverage = {
        "previous_source_kind": previous_source.get("source_kind"),
        "current_source_kind": current_source.get("source_kind"),
        "previous_read_scope": previous_source.get("read_scope"),
        "current_read_scope": current_source.get("read_scope"),
        "previous_truncated": previous_source.get("truncated"),
        "current_truncated": current_source.get("truncated"),
        "limitations": limitations,
    }
    if existing is None:
        existing = FileVersionDiff(
            id=new_id("file_version_diff"),
            storage_object_id=storage_object.id,
            previous_version_id=previous_version.id,
            previous_sha256_hex=previous_version.sha256_hex,
            current_sha256_hex=current_sha256_hex,
            diff_kind=diff_kind,
            summary_text=summary_text,
            added_lines_json=json.dumps(added_lines, ensure_ascii=True),
            removed_lines_json=json.dumps(removed_lines, ensure_ascii=True),
            coverage_json=json.dumps(coverage, ensure_ascii=True, sort_keys=True),
            created_at=now,
            updated_at=now,
            version=1,
        )
        session.add(existing)
    else:
        existing.previous_sha256_hex = previous_version.sha256_hex
        existing.diff_kind = diff_kind
        existing.summary_text = summary_text
        existing.added_lines_json = json.dumps(added_lines, ensure_ascii=True)
        existing.removed_lines_json = json.dumps(removed_lines, ensure_ascii=True)
        existing.coverage_json = json.dumps(coverage, ensure_ascii=True, sort_keys=True)
        existing.updated_at = now
        existing.version += 1
    return existing


def latest_file_version_diff(
    session: DatabaseSession,
    *,
    storage_object_id: str,
    storage_object_version_id: str | None,
    source_sha256_hex: str,
) -> FileVersionDiff | None:
    if storage_object_version_id is None:
        query = (
            select(FileVersionDiff)
            .where(FileVersionDiff.storage_object_id == storage_object_id)
            .where(FileVersionDiff.current_sha256_hex == source_sha256_hex)
        )
        return session.scalar(
            query.order_by(
                FileVersionDiff.updated_at.desc(),
                FileVersionDiff.created_at.desc(),
            ).limit(1)
        )

    selected_version = session.get(StorageObjectVersion, storage_object_version_id)
    if (
        selected_version is None
        or selected_version.storage_object_id != storage_object_id
        or selected_version.version_number <= 1
    ):
        return None
    previous_version = file_version_for_number(
        session,
        storage_object_id=storage_object_id,
        version_number=selected_version.version_number - 1,
    )
    if previous_version is None:
        return None
    return session.scalar(
        select(FileVersionDiff)
        .where(FileVersionDiff.storage_object_id == storage_object_id)
        .where(FileVersionDiff.previous_version_id == previous_version.id)
        .where(FileVersionDiff.current_sha256_hex == source_sha256_hex)
        .order_by(FileVersionDiff.updated_at.desc(), FileVersionDiff.created_at.desc())
        .limit(1)
    )


def storage_version_by_sha256(
    session: DatabaseSession,
    *,
    storage_object_id: str,
    sha256_hex: str,
) -> StorageObjectVersion | None:
    return session.scalar(
        select(StorageObjectVersion)
        .where(StorageObjectVersion.storage_object_id == storage_object_id)
        .where(StorageObjectVersion.sha256_hex == sha256_hex)
        .order_by(StorageObjectVersion.version_number.desc())
        .limit(1)
    )


def storage_source_for_sha256(
    session: DatabaseSession,
    *,
    storage_object: StorageObject,
    sha256_hex: str,
) -> dict[str, object] | None:
    if storage_object.sha256_hex == sha256_hex:
        return storage_summary_source(storage_object, None, session)
    version = storage_version_by_sha256(
        session,
        storage_object_id=storage_object.id,
        sha256_hex=sha256_hex,
    )
    if version is None:
        return None
    return storage_summary_source(storage_object, version, session)


def file_version_diff_display_lines(
    session: DatabaseSession,
    *,
    storage_object: StorageObject,
    diff: FileVersionDiff,
) -> list[dict[str, str]]:
    previous_version = session.get(StorageObjectVersion, diff.previous_version_id)
    if previous_version is None or previous_version.storage_object_id != storage_object.id:
        return []
    current_source = storage_source_for_sha256(
        session,
        storage_object=storage_object,
        sha256_hex=diff.current_sha256_hex,
    )
    if current_source is None:
        return []
    try:
        previous_source = storage_summary_source(storage_object, previous_version, session)
    except Exception:
        return []
    previous_text = str(previous_source.get("source_text") or "")
    current_text = str(current_source.get("source_text") or "")
    if previous_text.strip() == "" and current_text.strip() == "":
        return []

    raw_lines: list[dict[str, str]] = []
    changed_indexes: list[int] = []
    context_radius = 3
    edge_context_lines = 3
    display_lines: list[dict[str, str]] = []
    for line in difflib.ndiff(previous_text.splitlines(), current_text.splitlines()):
        if line.startswith("? "):
            continue
        marker = line[:2]
        if marker == "+ ":
            kind = "added"
            text_value = line[2:]
        elif marker == "- ":
            kind = "removed"
            text_value = line[2:]
        else:
            kind = "context"
            text_value = line[2:] if marker == "  " else line
        raw_lines.append({"kind": kind, "text": text_value[:500]})
        if kind != "context":
            changed_indexes.append(len(raw_lines) - 1)

    visible_indexes: set[int] = set()
    visible_indexes.update(range(min(edge_context_lines, len(raw_lines))))
    visible_indexes.update(range(max(0, len(raw_lines) - edge_context_lines), len(raw_lines)))
    for index in changed_indexes:
        start = max(0, index - context_radius)
        end = min(len(raw_lines), index + context_radius + 1)
        visible_indexes.update(range(start, end))

    last_index = -1
    for index in sorted(visible_indexes):
        if last_index != -1 and index > last_index + 1:
            hidden_count = index - last_index - 1
            display_lines.append(
                {
                    "kind": "ellipsis",
                    "text": f"... {hidden_count} unchanged line(s) omitted ...",
                }
            )
        display_lines.append(raw_lines[index])
        last_index = index
        if len(display_lines) >= 300:
            display_lines.append({"kind": "ellipsis", "text": "... diff truncated ..."})
            break
    return display_lines


def target_storage_version_number(
    session: DatabaseSession,
    *,
    storage_object_id: str,
    version: StorageObjectVersion | None,
) -> int:
    if version is not None:
        return int(version.version_number)
    current = session.scalar(
        select(func.max(StorageObjectVersion.version_number)).where(
            StorageObjectVersion.storage_object_id == storage_object_id
        )
    )
    return int(current or 0) + 1


def latest_incremental_base_summary(
    session: DatabaseSession,
    *,
    storage_object_id: str,
    target_version_number: int,
) -> tuple[FileSummary, StorageObjectVersion] | None:
    row = session.execute(
        select(FileSummary, StorageObjectVersion)
        .join(
            StorageObjectVersion,
            FileSummary.storage_object_version_id == StorageObjectVersion.id,
        )
        .where(FileSummary.storage_object_id == storage_object_id)
        .where(FileSummary.summary_type == "llm_digest")
        .where(StorageObjectVersion.storage_object_id == storage_object_id)
        .where(StorageObjectVersion.version_number < target_version_number)
        .order_by(
            StorageObjectVersion.version_number.desc(),
            FileSummary.updated_at.desc(),
            FileSummary.created_at.desc(),
        )
        .limit(1)
    ).first()
    if row is None:
        return None
    return row[0], row[1]


def file_version_for_number(
    session: DatabaseSession,
    *,
    storage_object_id: str,
    version_number: int,
) -> StorageObjectVersion | None:
    return session.scalar(
        select(StorageObjectVersion)
        .where(StorageObjectVersion.storage_object_id == storage_object_id)
        .where(StorageObjectVersion.version_number == version_number)
        .limit(1)
    )


def incremental_diff_chain(
    session: DatabaseSession,
    *,
    storage_object: StorageObject,
    base_version: StorageObjectVersion,
    target_version: StorageObjectVersion | None,
) -> list[FileVersionDiff]:
    target_number = target_storage_version_number(
        session,
        storage_object_id=storage_object.id,
        version=target_version,
    )
    diffs: list[FileVersionDiff] = []
    for version_number in range(base_version.version_number, target_number):
        previous_version = file_version_for_number(
            session,
            storage_object_id=storage_object.id,
            version_number=version_number,
        )
        if previous_version is None:
            return []
        if version_number + 1 == target_number and target_version is None:
            next_sha256_hex = storage_object.sha256_hex
        else:
            next_version = file_version_for_number(
                session,
                storage_object_id=storage_object.id,
                version_number=version_number + 1,
            )
            if next_version is None:
                return []
            next_sha256_hex = next_version.sha256_hex
        diff = session.scalar(
            select(FileVersionDiff)
            .where(FileVersionDiff.storage_object_id == storage_object.id)
            .where(FileVersionDiff.previous_version_id == previous_version.id)
            .where(FileVersionDiff.current_sha256_hex == next_sha256_hex)
            .order_by(FileVersionDiff.updated_at.desc(), FileVersionDiff.created_at.desc())
            .limit(1)
        )
        if diff is None:
            return []
        diffs.append(diff)
    return diffs


def file_summary_incremental_payload(
    *,
    base_summary: FileSummary,
    base_version: StorageObjectVersion,
    diffs: list[FileVersionDiff],
) -> dict[str, object]:
    return {
        "base_summary_id": base_summary.id,
        "base_storage_object_version_id": base_version.id,
        "base_version_number": base_version.version_number,
        "base_source_sha256_hex": base_summary.source_sha256_hex,
        "base_file_description": base_summary.file_description,
        "base_summary_points": safe_json_list(base_summary.summary_points_json)[:5],
        "base_llm_digest": base_summary.llm_digest,
        "diffs": [
            {
                "previous_version_id": diff.previous_version_id,
                "previous_sha256_hex": diff.previous_sha256_hex,
                "current_sha256_hex": diff.current_sha256_hex,
                "diff_kind": diff.diff_kind,
                "summary_text": diff.summary_text,
                "added_lines": safe_json_list(diff.added_lines_json),
                "removed_lines": safe_json_list(diff.removed_lines_json),
                "coverage": safe_json_dict(diff.coverage_json),
            }
            for diff in diffs
        ],
    }


def normalize_summary_points(value: object) -> list[str]:
    if not isinstance(value, list):
        return []
    points = []
    for item in value:
        text_value = str(item).strip()
        if text_value:
            points.append(text_value)
    return points[:5]


def int_or_none(value: object) -> int | None:
    return value if isinstance(value, int) else None


def ensure_storage_object_llm_digest(
    session: DatabaseSession,
    *,
    storage_object: StorageObject,
    version: StorageObjectVersion | None,
    force: bool = False,
) -> FileSummary:
    if not bool(storage_object.llm_input_allowed):
        raise json_error(409, "LLM_INPUT_BLOCKED", "LLM input is blocked for this file.")

    source = storage_summary_source(storage_object, version, session)
    storage_object_version_id = (
        source["storage_object_version_id"]
        if isinstance(source["storage_object_version_id"], str)
        else None
    )
    existing_summary = latest_matching_file_summary(
        session,
        storage_object_id=storage_object.id,
        storage_object_version_id=storage_object_version_id,
        source_sha256_hex=str(source["sha256_hex"]),
    )
    if existing_summary is not None and not force:
        return existing_summary

    target_version_number = target_storage_version_number(
        session,
        storage_object_id=storage_object.id,
        version=version,
    )
    incremental_payload: dict[str, object] | None = None
    incremental_base = (
        latest_incremental_base_summary(
            session,
            storage_object_id=storage_object.id,
            target_version_number=target_version_number,
        )
        if existing_summary is None
        else None
    )
    if incremental_base is not None:
        base_summary, base_version = incremental_base
        diff_chain = incremental_diff_chain(
            session,
            storage_object=storage_object,
            base_version=base_version,
            target_version=version,
        )
        if diff_chain:
            incremental_payload = file_summary_incremental_payload(
                base_summary=base_summary,
                base_version=base_version,
                diffs=diff_chain,
            )

    provider = build_file_summary_provider()
    now = jst_iso()
    input_payload = {
        "generation_mode": "incremental_from_digest"
        if incremental_payload is not None
        else "full_source",
        "storage_object_id": storage_object.id,
        "storage_object_version_id": source["storage_object_version_id"],
        "filename": source["filename"],
        "content_type": source["content_type"],
        "byte_size": source["byte_size"],
        "sha256_hex": source["sha256_hex"],
        "source_kind": source["source_kind"],
        "read_scope": source["read_scope"],
        "truncated": source["truncated"],
        "limitations": source["limitations"],
        "source_text": "" if incremental_payload is not None else source["source_text"],
        "incremental_source": incremental_payload,
    }
    try:
        provider_response = provider.complete_json(
            function_type=FUNCTION_TYPE_FILE_SUMMARY,
            input_payload=input_payload,
        )
    except OpenAIProviderError as error:
        raise json_error(
            502,
            "LLM_PROVIDER_ERROR",
            f"File digest generation failed: {error}",
        ) from error

    output = provider_response.output
    description = str(output.get("file_description") or "").strip()
    if description == "":
        raise json_error(502, "LLM_PROVIDER_ERROR", "File digest is missing a description.")
    summary_points = normalize_summary_points(output.get("summary_points"))
    llm_digest = str(output.get("llm_digest") or "").strip()
    structured_digest = (
        output.get("structured_digest")
        if isinstance(output.get("structured_digest"), dict)
        else {}
    )
    coverage = output.get("coverage") if isinstance(output.get("coverage"), dict) else {}

    llm_run = LlmRun(
        id=new_id("llm_run"),
        function_type=FUNCTION_TYPE_FILE_SUMMARY,
        provider_name=provider.provider_name,
        model_name=provider.model_name,
        prompt_version_id=None,
        input_hash=None,
        input_source_json=json.dumps(
            {
                "storage_object_id": storage_object.id,
                "storage_object_version_id": source["storage_object_version_id"],
                "sha256_hex": source["sha256_hex"],
                "filename": source["filename"],
                "content_type": source["content_type"],
            },
            ensure_ascii=True,
            sort_keys=True,
        ),
        input_diagnostic_json=json.dumps(
            {
                "source_kind": source["source_kind"],
                "read_scope": source["read_scope"],
                "truncated": source["truncated"],
                "source_text_length": len(str(source["source_text"])),
                "limitations": source["limitations"],
                "generation_mode": input_payload["generation_mode"],
                "incremental_diff_count": len(
                    incremental_payload.get("diffs", [])
                    if incremental_payload is not None
                    else []
                ),
            },
            ensure_ascii=True,
            sort_keys=True,
        ),
        applied_instruction_rule_ids_json=json.dumps([], ensure_ascii=True),
        output_json=json.dumps(output, ensure_ascii=True, sort_keys=True),
        output_text_preview=provider_response.output_preview,
        status="succeeded",
        error_type=None,
        error_message=None,
        retry_count=0,
        max_retry_count=3,
        prompt_tokens=provider_response.prompt_tokens,
        completion_tokens=provider_response.completion_tokens,
        total_tokens=provider_response.total_tokens,
        estimated_cost=provider_response.estimated_cost,
        started_at=now,
        finished_at=now,
        created_at=now,
    )
    session.add(llm_run)

    summary = existing_summary
    if summary is None:
        summary = FileSummary(
            id=new_id("file_summary"),
            storage_object_id=storage_object.id,
            storage_object_version_id=storage_object_version_id,
            source_sha256_hex=str(source["sha256_hex"]),
            source_filename=(
                str(source["filename"]) if source["filename"] is not None else None
            ),
            source_content_type=(
                str(source["content_type"])
                if source["content_type"] is not None
                else None
            ),
            source_byte_size=int(source["byte_size"] or 0),
            summary_type="llm_digest",
            file_description=description,
            summary_points_json=json.dumps(
                summary_points,
                ensure_ascii=True,
                sort_keys=True,
            ),
            llm_digest=llm_digest,
            structured_digest_json=json.dumps(
                structured_digest,
                ensure_ascii=True,
                sort_keys=True,
            ),
            coverage_json=json.dumps(coverage, ensure_ascii=True, sort_keys=True),
            token_estimate=int_or_none(output.get("token_estimate")),
            llm_run_id=llm_run.id,
            created_at=now,
            updated_at=now,
            version=1,
        )
        session.add(summary)
    else:
        summary.source_filename = (
            str(source["filename"]) if source["filename"] is not None else None
        )
        summary.source_content_type = (
            str(source["content_type"]) if source["content_type"] is not None else None
        )
        summary.source_byte_size = int(source["byte_size"] or 0)
        summary.file_description = description
        summary.summary_points_json = json.dumps(
            summary_points,
            ensure_ascii=True,
            sort_keys=True,
        )
        summary.llm_digest = llm_digest
        summary.structured_digest_json = json.dumps(
            structured_digest,
            ensure_ascii=True,
            sort_keys=True,
        )
        summary.coverage_json = json.dumps(coverage, ensure_ascii=True, sort_keys=True)
        summary.token_estimate = int_or_none(output.get("token_estimate"))
        summary.llm_run_id = llm_run.id
        summary.updated_at = now
        summary.version += 1

    storage_object.updated_at = now
    storage_object.version += 1
    record_storage_operation(
        session,
        operation_type="llm_digest_prepared",
        now=now,
        storage_object=storage_object,
        details={
            "file_summary_id": summary.id,
            "storage_object_version_id": source["storage_object_version_id"],
            "source_sha256_hex": source["sha256_hex"],
            "generation_mode": input_payload["generation_mode"],
            "incremental_base_summary_id": (
                incremental_payload.get("base_summary_id")
                if incremental_payload is not None
                else None
            ),
        },
    )
    return summary


@router.get("/locations")
def list_storage_locations(
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    aggregate_rows = session.execute(
        select(
            StorageObject.location_id,
            func.count(),
            func.coalesce(func.sum(StorageObject.byte_size), 0),
        )
        .where(StorageObject.status == "active")
        .where(StorageObject.scope == "managed")
        .group_by(StorageObject.location_id)
    ).all()
    aggregates = {
        row[0]: {
            "object_count": int(row[1] or 0),
            "active_byte_size": int(row[2] or 0),
        }
        for row in aggregate_rows
    }
    locations = session.scalars(
        select(StorageLocation).order_by(StorageLocation.kind, StorageLocation.label)
    ).all()
    return {
        "ok": True,
        "data": {
            "items": [
                storage_location_data(
                    location,
                    object_count=aggregates.get(location.id, {}).get("object_count", 0),
                    active_byte_size=aggregates.get(location.id, {}).get(
                        "active_byte_size",
                        0,
                    ),
                )
                for location in locations
            ]
        },
    }


@router.get("/file-icons")
def list_file_icon_settings(
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    settings = session.scalars(
        select(FileIconSetting).order_by(
            FileIconSetting.created_at.asc(),
            FileIconSetting.id.asc(),
        )
    ).all()
    return {"ok": True, "data": {"items": [file_icon_setting_data(item) for item in settings]}}


@router.post("/file-icons")
def create_file_icon_setting(
    payload: FileIconSettingCreate,
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    extensions = normalize_file_icon_extensions(payload.extensions)
    now = jst_iso()
    icon_object = save_file_icon_storage_object(
        session,
        filename=normalized_optional_filename(payload.icon_filename),
        content_type=payload.icon_content_type,
        data_base64=payload.icon_data_base64,
        now=now,
    )
    setting = FileIconSetting(
        id=new_id("file_icon_setting"),
        icon_filename=normalized_optional_filename(payload.icon_filename),
        storage_object_id=icon_object.id,
        icon_content_type=icon_object.content_type or payload.icon_content_type.strip().lower(),
        icon_data_url="",
        extensions_json=json.dumps(extensions, ensure_ascii=True),
        created_at=now,
        updated_at=now,
        version=1,
    )
    session.add(setting)
    session.commit()
    return {"ok": True, "data": {"file_icon": file_icon_setting_data(setting)}}


@router.patch("/file-icons/{file_icon_id}")
def update_file_icon_setting(
    file_icon_id: str,
    payload: FileIconSettingUpdate,
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    setting = session.get(FileIconSetting, file_icon_id)
    if setting is None:
        raise json_error(404, "NOT_FOUND", "File icon setting not found.")
    now = jst_iso()
    previous_storage_object_id: str | None = None
    if payload.extensions is not None:
        setting.extensions_json = json.dumps(
            normalize_file_icon_extensions(payload.extensions),
            ensure_ascii=True,
        )
    if payload.icon_content_type is not None and payload.icon_data_base64 is not None:
        previous_storage_object_id = setting.storage_object_id
        icon_object = save_file_icon_storage_object(
            session,
            filename=normalized_optional_filename(payload.icon_filename),
            content_type=payload.icon_content_type,
            data_base64=payload.icon_data_base64,
            now=now,
        )
        setting.storage_object_id = icon_object.id
        setting.icon_content_type = (
            icon_object.content_type or payload.icon_content_type.strip().lower()
        )
        setting.icon_data_url = ""
        setting.icon_filename = normalized_optional_filename(payload.icon_filename)
    elif payload.icon_content_type is not None or payload.icon_data_base64 is not None:
        raise json_error(
            422,
            "VALIDATION_ERROR",
            "Icon content type and image data must be updated together.",
        )
    setting.updated_at = now
    setting.version += 1
    delete_file_icon_storage_object_if_unused(
        session,
        storage_object_id=previous_storage_object_id,
        setting_id=setting.id,
        now=now,
    )
    session.commit()
    return {"ok": True, "data": {"file_icon": file_icon_setting_data(setting)}}


@router.delete("/file-icons/{file_icon_id}")
def delete_file_icon_setting(
    file_icon_id: str,
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    setting = session.get(FileIconSetting, file_icon_id)
    if setting is None:
        raise json_error(404, "NOT_FOUND", "File icon setting not found.")
    storage_object_id = setting.storage_object_id
    now = jst_iso()
    session.delete(setting)
    session.flush()
    delete_file_icon_storage_object_if_unused(
        session,
        storage_object_id=storage_object_id,
        setting_id=file_icon_id,
        now=now,
    )
    session.commit()
    return {"ok": True, "data": {"deleted": True}}


@router.get("/directories")
def list_storage_directories(
    parent_id: str = "root",
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    normalized_parent_id = normalize_directory_id(parent_id)
    ensure_storage_directory(session, normalized_parent_id)
    statement = (
        select(StorageDirectory)
        .where(StorageDirectory.status == "active")
        .where(StorageDirectory.directory_kind.in_(ACTIVE_STORAGE_DIRECTORY_KINDS))
        .where(StorageDirectory.id.not_in(HIDDEN_STORAGE_DIRECTORY_IDS))
        .order_by(StorageDirectory.name, StorageDirectory.id)
    )
    if normalized_parent_id is None:
        statement = statement.where(StorageDirectory.parent_id.is_(None))
    else:
        statement = statement.where(StorageDirectory.parent_id == normalized_parent_id)
    directories = session.scalars(statement).all()
    return {
        "ok": True,
        "data": {
            "items": [storage_directory_data(directory) for directory in directories],
            "breadcrumbs": storage_directory_breadcrumbs(session, normalized_parent_id),
        },
    }


@router.post("/directories")
def create_storage_directory(
    payload: StorageDirectoryCreate,
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    name = payload.name.strip()
    if name == "":
        raise json_error(422, "VALIDATION_ERROR", "Directory name is required.")
    if "/" in name or "\\" in name:
        raise json_error(422, "VALIDATION_ERROR", "Directory name cannot contain slashes.")
    parent_id = normalize_directory_id(payload.parent_id)
    ensure_storage_directory(session, parent_id)
    duplicate_statement = (
        select(StorageDirectory)
        .where(StorageDirectory.status == "active")
        .where(StorageDirectory.name == name)
    )
    if parent_id is None:
        duplicate_statement = duplicate_statement.where(StorageDirectory.parent_id.is_(None))
    else:
        duplicate_statement = duplicate_statement.where(StorageDirectory.parent_id == parent_id)
    if session.scalar(duplicate_statement.limit(1)) is not None:
        raise json_error(409, "DUPLICATE_DIRECTORY", "Directory already exists.")
    now = jst_iso()
    directory = StorageDirectory(
        id=new_id("storage_directory"),
        parent_id=parent_id,
        directory_kind="normal",
        case_id=None,
        name=name,
        status="active",
        created_at=now,
        updated_at=now,
        version=1,
    )
    session.add(directory)
    session.commit()
    return {"ok": True, "data": {"directory": storage_directory_data(directory)}}


@router.patch("/directories/{directory_id}/parent")
def update_storage_directory_parent(
    directory_id: str,
    payload: StorageDirectoryParentPatch,
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    directory = ensure_storage_directory(session, directory_id)
    if directory is None:
        raise json_error(404, "NOT_FOUND", "Storage directory not found.")
    if directory.directory_kind != "normal":
        raise json_error(
            409,
            "DIRECTORY_MOVE_PROTECTED",
            "Only normal directories can be moved.",
        )
    if (
        directory.case_id is not None
        and directory.id == case_handover_storage_directory_id(directory.case_id)
    ):
        raise json_error(
            409,
            "DIRECTORY_MOVE_PROTECTED",
            "Case handover directory cannot be moved.",
        )

    parent_id = normalize_directory_id(payload.parent_id)
    ensure_storage_directory(session, parent_id)
    if parent_id == directory.id:
        raise json_error(
            409,
            "INVALID_DIRECTORY_MOVE",
            "Directory cannot be moved into itself.",
        )
    if parent_id is not None and parent_id in storage_directory_descendant_ids(
        session,
        directory.id,
    ):
        raise json_error(
            409,
            "INVALID_DIRECTORY_MOVE",
            "Directory cannot be moved into its descendant.",
        )

    duplicate_statement = (
        select(StorageDirectory)
        .where(StorageDirectory.status == "active")
        .where(StorageDirectory.name == directory.name)
        .where(StorageDirectory.id != directory.id)
    )
    if parent_id is None:
        duplicate_statement = duplicate_statement.where(StorageDirectory.parent_id.is_(None))
    else:
        duplicate_statement = duplicate_statement.where(StorageDirectory.parent_id == parent_id)
    if session.scalar(duplicate_statement.limit(1)) is not None:
        raise json_error(409, "DUPLICATE_DIRECTORY", "Directory already exists.")

    if directory.parent_id != parent_id:
        directory.parent_id = parent_id
        directory.updated_at = jst_iso()
        directory.version += 1
        session.commit()

    return {"ok": True, "data": {"directory": storage_directory_data(directory)}}


@router.delete("/directories/{directory_id}")
def delete_storage_directory(
    directory_id: str,
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    directory = ensure_storage_directory(session, directory_id)
    if directory is None:
        raise json_error(404, "NOT_FOUND", "Storage directory not found.")
    if directory.directory_kind == "case" and directory.case_id is not None:
        case = session.get(Case, directory.case_id)
        if case is not None:
            raise json_error(
                409,
                "CASE_DIRECTORY_PROTECTED",
                "Case directory cannot be deleted while the Case exists.",
            )
    if directory.directory_kind == "task":
        task_id = session.scalar(
            select(Task.id)
            .where(Task.storage_directory_id == directory.id)
            .where(Task.deleted_at.is_(None))
            .limit(1)
        )
        if task_id is not None:
            raise json_error(
                409,
                "TASK_DIRECTORY_PROTECTED",
                "Task directory cannot be deleted while the Task exists.",
            )

    to_visit = [directory]
    directories: list[StorageDirectory] = []
    seen: set[str] = set()
    while to_visit:
        current = to_visit.pop()
        if current.id in seen:
            continue
        seen.add(current.id)
        directories.append(current)
        children = session.scalars(
            select(StorageDirectory)
            .where(StorageDirectory.status == "active")
            .where(StorageDirectory.parent_id == current.id)
        ).all()
        to_visit.extend(children)

    directory_ids = [item.id for item in directories]
    storage_objects = session.scalars(
        select(StorageObject)
        .where(StorageObject.status == "active")
        .where(StorageObject.scope == "managed")
        .where(StorageObject.directory_id.in_(directory_ids))
    ).all()
    now = jst_iso()
    restored_count = 0
    for storage_object in storage_objects:
        restored = delete_managed_storage_object_for_request(
            storage_object,
            session=session,
            now=now,
        )
        if restored is not None:
            restored_count += 1
    for current in directories:
        current.status = "deleted"
        current.updated_at = now
        current.version += 1
    session.commit()
    return {
        "ok": True,
        "data": {
            "deleted_directory_id": directory.id,
            "deleted_directory_count": len(directories),
            "deleted_object_count": len(storage_objects),
            "restored_attachment_count": restored_count,
        },
    }


@router.get("/objects")
def list_storage_objects(
    status: str = "active",
    location_id: str = "all",
    directory_id: str = "root",
    limit: int = 100,
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    safe_limit = max(1, min(limit, 500))
    normalized_directory_id = normalize_directory_id(directory_id)
    ensure_storage_directory(session, normalized_directory_id)
    statement = select(StorageObject).order_by(
        StorageObject.created_at.desc(),
        StorageObject.id.desc(),
    )
    statement = statement.where(StorageObject.scope == "managed")
    if normalized_directory_id is None:
        statement = statement.where(
            (StorageObject.directory_id.is_(None))
            | (StorageObject.directory_id.not_in(HIDDEN_STORAGE_DIRECTORY_IDS))
        )
    if status != "all":
        statement = statement.where(StorageObject.status == status)
    if location_id != "all":
        statement = statement.where(StorageObject.location_id == location_id)
    if normalized_directory_id is None:
        statement = statement.where(StorageObject.directory_id.is_(None))
    else:
        statement = statement.where(StorageObject.directory_id == normalized_directory_id)
    objects = session.scalars(statement.limit(safe_limit)).all()
    linked_object_ids: set[str] = set()
    case = storage_case_for_directory_id(session, normalized_directory_id)
    if case is not None:
        linked_objects = storage_linked_case_directory_objects(
            session,
            case,
            normalized_directory_id,
            status=status,
            location_id=location_id,
            exclude_ids={storage_object.id for storage_object in objects},
            limit=safe_limit,
        )
        linked_object_ids = {storage_object.id for storage_object in linked_objects}
        objects = sort_storage_objects_for_directory(
            [*objects, *linked_objects]
        )[:safe_limit]
    return {
        "ok": True,
        "data": {
            "items": [
                storage_object_data(
                    storage_object,
                    session,
                    display_source="link"
                    if storage_object.id in linked_object_ids
                    else "physical",
                )
                for storage_object in objects
            ]
        },
    }


@router.get("/search/objects")
def search_storage_objects(
    q: str = "",
    status: str = "active",
    directory_id: str = "root",
    recursive: bool = True,
    sort: str = "created_desc",
    extension: str = "all",
    limit: int = 200,
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    safe_limit = max(1, min(limit, 500))
    normalized_directory_id = normalize_directory_id(directory_id)
    ensure_storage_directory(session, normalized_directory_id)

    def base_statement():
        statement = (
            select(StorageObject)
            .outerjoin(GmailMessage, StorageObject.source_message_id == GmailMessage.id)
            .where(StorageObject.scope == "managed")
        )
        if status != "all":
            statement = statement.where(StorageObject.status == status)
        if recursive:
            if normalized_directory_id is not None:
                directory_ids = storage_directory_descendant_ids(
                    session,
                    normalized_directory_id,
                )
                statement = statement.where(StorageObject.directory_id.in_(directory_ids))
            else:
                statement = statement.where(
                    (StorageObject.directory_id.is_(None))
                    | (StorageObject.directory_id.not_in(HIDDEN_STORAGE_DIRECTORY_IDS))
                )
        elif normalized_directory_id is None:
            statement = statement.where(StorageObject.directory_id.is_(None))
        else:
            statement = statement.where(StorageObject.directory_id == normalized_directory_id)

        query_terms = [term for term in q.strip().lower().split() if term]
        for term in query_terms:
            pattern = f"%{term}%"
            statement = statement.where(
                or_(
                    func.lower(StorageObject.original_filename).like(pattern),
                    func.lower(GmailMessage.subject).like(pattern),
                )
            )
        return statement

    extension_statement = base_statement()
    extension_objects = session.scalars(extension_statement.limit(1000)).all()
    available_extensions = sorted(
        {
            item
            for item in (
                storage_file_extension(storage_object.original_filename)
                for storage_object in extension_objects
            )
            if item is not None
        }
    )

    statement = base_statement()
    normalized_extension = extension.strip().lower().lstrip(".")
    if normalized_extension not in {"", "all"}:
        statement = statement.where(
            func.lower(StorageObject.original_filename).like(f"%.{normalized_extension}")
        )
    if sort == "name":
        statement = statement.order_by(
            func.lower(StorageObject.original_filename).asc(),
            StorageObject.id.asc(),
        )
    elif sort == "created_asc":
        statement = statement.order_by(StorageObject.created_at.asc(), StorageObject.id.asc())
    else:
        statement = statement.order_by(StorageObject.created_at.desc(), StorageObject.id.desc())

    objects = session.scalars(statement.limit(safe_limit)).all()
    return {
        "ok": True,
        "data": {
            "items": [
                storage_object_data(storage_object, session) for storage_object in objects
            ],
            "extensions": available_extensions,
        },
    }


@router.post("/objects")
def upload_managed_object(
    payload: ManagedObjectUpload,
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    directory = ensure_storage_directory(session, payload.directory_id)
    data = decode_base64_payload(payload.data_base64)
    if len(data) == 0:
        raise json_error(422, "VALIDATION_ERROR", "Storage object is empty.")
    if len(data) > MAX_TMP_OBJECT_BYTES:
        raise json_error(413, "PAYLOAD_TOO_LARGE", "Storage object is too large.")

    content_type = (payload.content_type or "application/octet-stream").strip()
    if content_type == "":
        content_type = "application/octet-stream"
    storage_object = save_storage_object(
        session,
        scope="managed",
        filename=payload.filename,
        content_type=content_type,
        data=data,
        now=jst_iso(),
        directory_id=directory.id if directory is not None else None,
        source_type="direct_upload",
    )
    session.commit()
    return {
        "ok": True,
        "data": {"storage_object": storage_object_data(storage_object, session)},
    }


@router.post("/objects/upload")
async def upload_managed_object_multipart(
    file: UploadFile = File(...),
    directory_id: str | None = None,
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    directory = ensure_storage_directory(session, directory_id)
    storage_object = await save_uploaded_storage_object(
        session,
        scope="managed",
        upload=file,
        now=jst_iso(),
        directory_id=directory.id if directory is not None else None,
        source_type="direct_upload",
    )
    if (
        storage_object.directory_id == BOOKSHELF_STORAGE_DIRECTORY_ID
        and is_pdf_storage_file(storage_object.content_type, storage_object.original_filename)
    ):
        ensure_pdf_ocr_status(session, storage_object)
    session.commit()
    return {
        "ok": True,
        "data": {"storage_object": storage_object_data(storage_object, session)},
    }





async def read_bookshelf_upload_bytes(upload: UploadFile) -> bytes:
    data = await upload.read()
    await upload.close()
    if len(data) == 0:
        raise json_error(422, "VALIDATION_ERROR", "Storage object is empty.")
    if len(data) > MAX_TMP_OBJECT_BYTES:
        raise json_error(413, "PAYLOAD_TOO_LARGE", "Storage object is too large.")
    return data


def convert_epub_bytes_to_pdf_bytes(data: bytes, filename: str | None) -> bytes:
    command = shutil.which("ebook-convert")
    if command is None:
        raise json_error(503, "EPUB_CONVERTER_UNAVAILABLE", "ebook-convert is not installed.")
    with tempfile.TemporaryDirectory(prefix="caseclosed_epub_convert_") as tmpdir:
        tmp_path = Path(tmpdir)
        input_path = tmp_path / (Path(filename or "book.epub").name or "book.epub")
        if input_path.suffix.lower() != ".epub":
            input_path = input_path.with_suffix(".epub")
        output_path = tmp_path / "converted.pdf"
        input_path.write_bytes(data)
        try:
            result = subprocess.run(
                [command, str(input_path), str(output_path)],
                check=False,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=180,
            )
        except subprocess.TimeoutExpired as error:
            raise json_error(504, "EPUB_CONVERSION_TIMEOUT", "EPUB to PDF conversion timed out.") from error
        if result.returncode != 0:
            detail = (result.stderr or result.stdout).decode("utf-8", errors="replace").strip()
            message = "EPUB to PDF conversion failed."
            if detail:
                message = f"{message} {detail[-800:]}"
            raise json_error(422, "EPUB_CONVERSION_FAILED", message)
        if not output_path.is_file() or output_path.stat().st_size == 0:
            raise json_error(422, "EPUB_CONVERSION_FAILED", "EPUB to PDF conversion produced no PDF.")
        return output_path.read_bytes()


@router.get("/paper-journal-icons")
def list_paper_journal_icon_settings(
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    settings = session.scalars(
        select(PaperJournalIconSetting).order_by(
            PaperJournalIconSetting.match_journal.asc(),
            PaperJournalIconSetting.created_at.asc(),
            PaperJournalIconSetting.id.asc(),
        )
    ).all()
    return {"ok": True, "data": {"items": [paper_journal_icon_setting_data(item) for item in settings]}}


@router.post("/paper-journal-icons")
def create_paper_journal_icon_setting(
    payload: PaperJournalIconSettingCreate,
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    match_journal = normalize_paper_journal_match(payload.match_journal)
    if match_journal == "":
        raise json_error(422, "VALIDATION_ERROR", "Journal match text is required.")
    content_type, icon_data_url = paper_journal_icon_data_url(
        content_type=payload.icon_content_type,
        data_base64=payload.icon_data_base64,
    )
    now = jst_iso()
    setting = PaperJournalIconSetting(
        id=new_id("paper_journal_icon"),
        match_journal=match_journal,
        icon_filename=normalized_optional_filename(payload.icon_filename),
        icon_content_type=content_type,
        icon_data_url=icon_data_url,
        created_at=now,
        updated_at=now,
        version=1,
    )
    session.add(setting)
    session.commit()
    return {"ok": True, "data": {"journal_icon": paper_journal_icon_setting_data(setting)}}


@router.patch("/paper-journal-icons/{journal_icon_id}")
def update_paper_journal_icon_setting(
    journal_icon_id: str,
    payload: PaperJournalIconSettingUpdate,
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    setting = session.get(PaperJournalIconSetting, journal_icon_id)
    if setting is None:
        raise json_error(404, "NOT_FOUND", "Paper journal icon setting not found.")
    if payload.match_journal is not None:
        match_journal = normalize_paper_journal_match(payload.match_journal)
        if match_journal == "":
            raise json_error(422, "VALIDATION_ERROR", "Journal match text is required.")
        setting.match_journal = match_journal
    if payload.icon_content_type is not None and payload.icon_data_base64 is not None:
        content_type, icon_data_url = paper_journal_icon_data_url(
            content_type=payload.icon_content_type,
            data_base64=payload.icon_data_base64,
        )
        setting.icon_content_type = content_type
        setting.icon_data_url = icon_data_url
        setting.icon_filename = normalized_optional_filename(payload.icon_filename)
    elif payload.icon_content_type is not None or payload.icon_data_base64 is not None:
        raise json_error(
            422,
            "VALIDATION_ERROR",
            "Icon content type and image data must be updated together.",
        )
    setting.updated_at = jst_iso()
    setting.version += 1
    session.commit()
    return {"ok": True, "data": {"journal_icon": paper_journal_icon_setting_data(setting)}}


@router.delete("/paper-journal-icons/{journal_icon_id}")
def delete_paper_journal_icon_setting(
    journal_icon_id: str,
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    setting = session.get(PaperJournalIconSetting, journal_icon_id)
    if setting is None:
        raise json_error(404, "NOT_FOUND", "Paper journal icon setting not found.")
    session.delete(setting)
    session.commit()
    return {"ok": True, "data": {"deleted": True}}


@router.get("/papers")
def list_papers(
    status: str = "active",
    limit: int = 500,
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    safe_limit = max(1, min(limit, 500))
    statement = (
        select(Paper)
        .join(StorageObject, StorageObject.id == Paper.storage_object_id)
        .where(StorageObject.directory_id == PAPERS_STORAGE_DIRECTORY_ID)
        .order_by(Paper.title.asc(), Paper.id.asc())
        .limit(safe_limit)
    )
    if status != "all":
        statement = statement.where(Paper.status == status).where(StorageObject.status == status)
    papers = session.scalars(statement).all()
    items = [paper_data(paper, session) for paper in papers]
    if session.dirty or session.new:
        session.commit()
    return {"ok": True, "data": {"items": items}}


@router.get("/papers/{paper_id}")
def read_paper(
    paper_id: str,
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    paper = session.get(Paper, paper_id)
    if paper is None or paper.status != "active":
        raise json_error(404, "NOT_FOUND", "Paper not found.")
    storage_object = session.get(StorageObject, paper.storage_object_id)
    if (
        storage_object is None
        or storage_object.status != "active"
        or storage_object.directory_id != PAPERS_STORAGE_DIRECTORY_ID
    ):
        raise json_error(404, "NOT_FOUND", "Paper not found.")
    data = paper_data(paper, session)
    if session.dirty or session.new:
        session.commit()
    return {"ok": True, "data": {"paper": data}}


@router.post("/papers/upload")
async def upload_paper(
    file: UploadFile = File(...),
    bibtex_file: UploadFile = File(...),
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    if not is_pdf_filename(file.filename):
        raise json_error(422, "PAPER_PDF_ONLY", "PaperShelf accepts .pdf files only.")
    bibtex = await read_paper_bibtex_upload_text(bibtex_file)
    directory = ensure_storage_directory(session, PAPERS_STORAGE_DIRECTORY_ID)
    now = jst_iso()
    storage_object = await save_uploaded_storage_object(
        session,
        scope="managed",
        upload=file,
        now=now,
        directory_id=directory.id if directory is not None else None,
        source_type="paper",
    )
    paper = Paper(
        id=new_id("paper"),
        storage_object_id=storage_object.id,
        title=bibtex_field_value(bibtex, "title") or paper_title_from_filename(storage_object.original_filename),
        authors_text=bibtex_field_value(bibtex, "author") or "",
        bibtex=bibtex,
        summary="",
        tags_json="[]",
        status="active",
        created_at=now,
        updated_at=now,
        version=1,
    )
    session.add(paper)
    session.flush()
    upsert_paper_bibtex_entry(session, paper, now=now)
    session.commit()
    return {
        "ok": True,
        "data": {"paper": paper_data(paper, session)},
    }


@router.patch("/papers/{paper_id}")
def update_paper(
    paper_id: str,
    payload: PaperUpdate,
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    paper = session.get(Paper, paper_id)
    if paper is None or paper.status != "active":
        raise json_error(404, "NOT_FOUND", "Paper not found.")
    storage_object = session.get(StorageObject, paper.storage_object_id)
    if (
        storage_object is None
        or storage_object.status != "active"
        or storage_object.directory_id != PAPERS_STORAGE_DIRECTORY_ID
    ):
        raise json_error(404, "NOT_FOUND", "Paper not found.")
    if payload.title is not None:
        title = payload.title.strip()
        if title == "":
            raise json_error(422, "VALIDATION_ERROR", "Paper title is required.")
        paper.title = title
    if payload.authors_text is not None:
        paper.authors_text = payload.authors_text.strip()
    if payload.bibtex is not None:
        paper.bibtex = payload.bibtex.strip()
    if payload.summary is not None:
        paper.summary = payload.summary.strip()
    if payload.tags is not None:
        paper.tags_json = json.dumps(normalize_paper_tags(payload.tags), ensure_ascii=False)
    now = jst_iso()
    paper.updated_at = now
    paper.version += 1
    upsert_paper_bibtex_entry(session, paper, now=now)
    session.commit()
    return {"ok": True, "data": {"paper": paper_data(paper, session)}}


@router.delete("/papers/{paper_id}")
def delete_paper(
    paper_id: str,
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    paper = session.get(Paper, paper_id)
    if paper is None or paper.status != "active":
        raise json_error(404, "NOT_FOUND", "Paper not found.")
    storage_object = session.get(StorageObject, paper.storage_object_id)
    if (
        storage_object is None
        or storage_object.status != "active"
        or storage_object.directory_id != PAPERS_STORAGE_DIRECTORY_ID
        or storage_object.source_type != "paper"
        or not is_pdf_filename(storage_object.original_filename)
    ):
        raise json_error(404, "NOT_FOUND", "Paper not found.")
    result = purge_bookshelf_material_object(storage_object, session=session)
    purged_paper_id = paper.id
    session.execute(delete(PaperBibtexEntry).where(PaperBibtexEntry.paper_id == paper.id))
    session.delete(paper)
    session.commit()
    result["purged_paper_id"] = purged_paper_id
    return {"ok": True, "data": result}


@router.delete("/bookshelf/materials/{storage_object_id}")
def delete_bookshelf_material(
    storage_object_id: str,
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    storage_object = session.get(StorageObject, storage_object_id)
    if (
        storage_object is None
        or storage_object.status != "active"
        or storage_object.scope != "managed"
        or storage_object.directory_id != BOOKSHELF_STORAGE_DIRECTORY_ID
        or storage_object.source_type != "bookshelf_material"
        or not is_bookshelf_pdf_filename(storage_object.original_filename)
    ):
        raise json_error(404, "NOT_FOUND", "Bookshelf material not found.")
    result = purge_bookshelf_material_object(storage_object, session=session)
    session.commit()
    return {
        "ok": True,
        "data": result,
    }


@router.post("/bookshelf/materials/upload")
async def upload_bookshelf_material(
    file: UploadFile = File(...),
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    if not is_bookshelf_material_filename(file.filename):
        raise json_error(422, "BOOKSHELF_MATERIAL_ONLY", "Bookshelf accepts .pdf and .epub files only.")
    directory = ensure_storage_directory(session, BOOKSHELF_STORAGE_DIRECTORY_ID)
    if is_bookshelf_pdf_filename(file.filename):
        storage_object = await save_uploaded_storage_object(
            session,
            scope="managed",
            upload=file,
            now=jst_iso(),
            directory_id=directory.id if directory is not None else None,
            source_type="bookshelf_material",
        )
    else:
        epub_data = await read_bookshelf_upload_bytes(file)
        pdf_data = convert_epub_bytes_to_pdf_bytes(epub_data, file.filename)
        storage_object = save_storage_object(
            session,
            scope="managed",
            filename=converted_bookshelf_pdf_filename(file.filename),
            content_type="application/pdf",
            data=pdf_data,
            now=jst_iso(),
            directory_id=directory.id if directory is not None else None,
            source_type="bookshelf_material",
        )
    ensure_pdf_ocr_status(session, storage_object)
    session.commit()
    return {
        "ok": True,
        "data": {"storage_object": storage_object_data(storage_object, session)},
    }


@router.patch("/objects/{storage_object_id}/llm-input")
def update_storage_object_llm_input(
    storage_object_id: str,
    payload: StorageObjectLlmInputPatch,
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    storage_object = session.get(StorageObject, storage_object_id)
    if (
        storage_object is None
        or storage_object.status != "active"
        or storage_object.scope != "managed"
    ):
        raise json_error(404, "NOT_FOUND", "Storage object not found.")
    storage_object.llm_input_allowed = 1 if payload.llm_input_allowed else 0
    now = jst_iso()
    storage_object.updated_at = now
    storage_object.version += 1
    record_storage_operation(
        session,
        operation_type="llm_input_updated",
        now=now,
        storage_object=storage_object,
        details={"llm_input_allowed": payload.llm_input_allowed},
    )
    session.commit()
    return {
        "ok": True,
        "data": {"storage_object": storage_object_data(storage_object, session)},
    }


@router.patch("/objects/{storage_object_id}/filename")
def update_storage_object_filename(
    storage_object_id: str,
    payload: StorageObjectFilenamePatch,
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    storage_object = session.get(StorageObject, storage_object_id)
    if (
        storage_object is None
        or storage_object.status != "active"
        or storage_object.scope != "managed"
    ):
        raise json_error(404, "NOT_FOUND", "Storage object not found.")
    next_filename = normalized_optional_filename(payload.filename)
    if next_filename is None:
        raise json_error(422, "VALIDATION_ERROR", "Filename is required.")
    if len(next_filename) > 255 or any(character in next_filename for character in "\\/\r\n\t"):
        raise json_error(422, "VALIDATION_ERROR", "Filename contains invalid characters.")
    previous_filename = storage_object.original_filename
    if previous_filename == next_filename:
        return {
            "ok": True,
            "data": {"storage_object": storage_object_data(storage_object, session)},
        }
    now = jst_iso()
    storage_object.original_filename = next_filename
    storage_object.updated_at = now
    storage_object.version += 1
    record_storage_operation(
        session,
        operation_type="renamed",
        now=now,
        storage_object=storage_object,
        details={
            "old_filename": previous_filename,
            "new_filename": next_filename,
        },
    )
    session.commit()
    return {
        "ok": True,
        "data": {"storage_object": storage_object_data(storage_object, session)},
    }


@router.patch("/objects/{storage_object_id}/directory")
def update_storage_object_directory(
    storage_object_id: str,
    payload: StorageObjectDirectoryPatch,
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    storage_object = session.get(StorageObject, storage_object_id)
    if (
        storage_object is None
        or storage_object.status != "active"
        or storage_object.scope != "managed"
    ):
        raise json_error(404, "NOT_FOUND", "Storage object not found.")
    old_directory_id = storage_object.directory_id
    directory = ensure_storage_directory(session, payload.directory_id)
    new_directory_id = directory.id if directory is not None else None
    now = jst_iso()
    old_case_id = storage_directory_case_id(session, old_directory_id)
    target_case_id = storage_directory_case_id(session, new_directory_id)
    physical_target_case_id = storage_directory_case_id(session, storage_object.directory_id)
    if target_case_id is not None and physical_target_case_id != target_case_id:
        active_case_link = session.scalar(
            select(FileLink)
            .where(FileLink.storage_object_id == storage_object.id)
            .where(FileLink.linked_type == "case")
            .where(FileLink.linked_id == target_case_id)
            .where(FileLink.status == "active")
            .order_by(FileLink.created_at.asc(), FileLink.id.asc())
            .limit(1)
        )
        if active_case_link is not None:
            new_directory_id = active_case_link.directory_id or case_storage_directory_id(
                target_case_id
            )

    storage_object.directory_id = new_directory_id
    if is_case_handover_directory_descendant(session, storage_object.directory_id):
        storage_object.llm_input_allowed = 1
    moved_into_case_id = storage_directory_case_id(session, storage_object.directory_id)
    redundant_file_links: list[FileLink] = []
    if moved_into_case_id is not None:
        redundant_file_links = session.scalars(
            select(FileLink)
            .where(FileLink.storage_object_id == storage_object.id)
            .where(FileLink.linked_type == "case")
            .where(FileLink.linked_id == moved_into_case_id)
            .where(FileLink.status == "active")
        ).all()
        for file_link in redundant_file_links:
            file_link.status = "deleted"
            file_link.updated_at = now
            file_link.version += 1
    storage_object.updated_at = now
    storage_object.version += 1
    record_storage_operation(
        session,
        operation_type="moved",
        now=now,
        storage_object=storage_object,
        details={
            "old_directory_id": old_directory_id,
            "new_directory_id": storage_object.directory_id,
            "deleted_redundant_case_file_link_count": len(redundant_file_links),
            "old_case_id": old_case_id,
        },
    )
    session.commit()
    return {
        "ok": True,
        "data": {"storage_object": storage_object_data(storage_object, session)},
    }


@router.get("/objects/{storage_object_id}/linked-cases")
def list_storage_object_linked_cases(
    storage_object_id: str,
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    storage_object = session.get(StorageObject, storage_object_id)
    if (
        storage_object is None
        or storage_object.status != "active"
        or storage_object.scope != "managed"
    ):
        raise json_error(404, "NOT_FOUND", "Storage object not found.")
    return {
        "ok": True,
        "data": {
            "items": storage_object_linked_case_items(session, storage_object),
        },
    }


@router.post("/objects/{storage_object_id}/linked-cases")
def create_storage_object_case_link(
    storage_object_id: str,
    payload: StorageObjectCaseLinkCreate,
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    storage_object = session.get(StorageObject, storage_object_id)
    if (
        storage_object is None
        or storage_object.status != "active"
        or storage_object.scope != "managed"
    ):
        raise json_error(404, "NOT_FOUND", "Storage object not found.")
    case = session.get(Case, payload.case_id)
    if case is None:
        raise json_error(404, "NOT_FOUND", "Case not found.")

    now = jst_iso()
    physical_case_id = storage_directory_case_id(session, storage_object.directory_id)
    existing_link = session.scalar(
        select(FileLink)
        .where(FileLink.storage_object_id == storage_object.id)
        .where(FileLink.linked_type == "case")
        .where(FileLink.linked_id == case.id)
        .order_by(FileLink.created_at.asc(), FileLink.id.asc())
        .limit(1)
    )
    link_changed = False
    if physical_case_id != case.id:
        if existing_link is None:
            existing_link = FileLink(
                id=f"file_link_{uuid4().hex}",
                storage_object_id=storage_object.id,
                linked_type="case",
                linked_id=case.id,
                label=None,
                status="active",
                created_at=now,
                updated_at=now,
                version=1,
            )
            session.add(existing_link)
            link_changed = True
        elif existing_link.status != "active":
            existing_link.status = "active"
            existing_link.updated_at = now
            existing_link.version += 1
            link_changed = True

    if link_changed:
        case.updated_at = now
        case.version += 1
        record_storage_operation(
            session,
            operation_type="case_linked",
            now=now,
            storage_object=storage_object,
            details={"case_id": case.id, "case_name": case.name},
        )
        session.commit()

    return {
        "ok": True,
        "data": {
            "items": storage_object_linked_case_items(session, storage_object),
        },
    }


@router.delete("/objects/{storage_object_id}/linked-cases/{case_id}")
def delete_storage_object_case_link(
    storage_object_id: str,
    case_id: str,
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    storage_object = session.get(StorageObject, storage_object_id)
    if (
        storage_object is None
        or storage_object.status != "active"
        or storage_object.scope != "managed"
    ):
        raise json_error(404, "NOT_FOUND", "Storage object not found.")
    case = session.get(Case, case_id)
    if case is None:
        raise json_error(404, "NOT_FOUND", "Case not found.")

    now = jst_iso()
    active_links = session.scalars(
        select(FileLink)
        .where(FileLink.storage_object_id == storage_object.id)
        .where(FileLink.linked_type == "case")
        .where(FileLink.linked_id == case.id)
        .where(FileLink.status == "active")
    ).all()
    for file_link in active_links:
        file_link.status = "deleted"
        file_link.updated_at = now
        file_link.version += 1

    if active_links:
        case.updated_at = now
        case.version += 1
        record_storage_operation(
            session,
            operation_type="case_unlinked",
            now=now,
            storage_object=storage_object,
            details={
                "case_id": case.id,
                "case_name": case.name,
                "deleted_file_link_count": len(active_links),
            },
        )
        session.commit()

    return {
        "ok": True,
        "data": {
            "items": storage_object_linked_case_items(session, storage_object),
        },
    }


@router.delete("/objects/{storage_object_id}")
def delete_managed_storage_object(
    storage_object_id: str,
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    storage_object = session.get(StorageObject, storage_object_id)
    if (
        storage_object is None
        or storage_object.status != "active"
        or storage_object.scope != "managed"
    ):
        raise json_error(404, "NOT_FOUND", "Storage object not found.")

    now = jst_iso()
    restored_storage_object = delete_managed_storage_object_for_request(
        storage_object,
        session=session,
        now=now,
    )
    session.commit()
    return {
        "ok": True,
        "data": {
            "deleted_storage_object_id": storage_object.id,
            "restored_storage_object": (
                storage_object_data(restored_storage_object, session)
                if restored_storage_object is not None
                else None
            ),
            "source_type": storage_object.source_type or "direct_upload",
        },
    }


@router.get("/objects/{storage_object_id}/pdf-ocr/status")
def read_storage_object_pdf_ocr_status(
    storage_object_id: str,
    refresh: bool = Query(False),
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    storage_object = require_pdf_storage_object_for_ocr(session, storage_object_id)
    status_row = ensure_pdf_ocr_status(session, storage_object, refresh=refresh)
    if session.dirty or session.new:
        session.commit()
    return {
        "ok": True,
        "data": {
            "ocr_status": pdf_ocr_status_data(status_row),
            "storage_object": storage_object_data(storage_object, session),
        },
    }


@router.post("/objects/{storage_object_id}/pdf-ocr/run")
def run_storage_object_pdf_ocr(
    storage_object_id: str,
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    storage_object = require_pdf_storage_object_for_ocr(session, storage_object_id)
    now = jst_iso()
    version, status_row = run_ocrmypdf_for_storage_object(session, storage_object, now=now)
    session.commit()
    return {
        "ok": True,
        "data": {
            "storage_object": storage_object_data(storage_object, session),
            "ocr_status": pdf_ocr_status_data(status_row),
            "version": storage_object_version_data(version) if version is not None else None,
            "skipped": version is None,
        },
    }


@router.get("/objects/{storage_object_id}/versions")
def list_storage_object_versions(
    storage_object_id: str,
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    storage_object = session.get(StorageObject, storage_object_id)
    if (
        storage_object is None
        or storage_object.status != "active"
        or storage_object.scope != "managed"
    ):
        raise json_error(404, "NOT_FOUND", "Storage object not found.")
    versions = session.scalars(
        select(StorageObjectVersion)
        .where(StorageObjectVersion.storage_object_id == storage_object.id)
        .order_by(
            StorageObjectVersion.version_number.desc(),
            StorageObjectVersion.id.desc(),
        )
    ).all()
    return {
        "ok": True,
        "data": {"items": [storage_object_version_data(version) for version in versions]},
    }


@router.post("/objects/{storage_object_id}/versions/upload")
async def upload_storage_object_version(
    storage_object_id: str,
    file: UploadFile = File(...),
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    storage_object = session.get(StorageObject, storage_object_id)
    if (
        storage_object is None
        or storage_object.status != "active"
        or storage_object.scope != "managed"
    ):
        raise json_error(404, "NOT_FOUND", "Storage object not found.")
    now = jst_iso()
    version = await update_storage_object_from_upload(
        session,
        storage_object,
        upload=file,
        now=now,
    )
    session.commit()
    return {
        "ok": True,
        "data": {
            "storage_object": storage_object_data(storage_object, session),
            "version": (
                storage_object_version_data(version) if version is not None else None
            ),
            "skipped": version is None,
            "skip_reason": "duplicate_content" if version is None else None,
        },
    }


def get_active_storage_object_version(
    session: DatabaseSession,
    storage_object_id: str,
    version_id: str,
) -> tuple[StorageObject, StorageObjectVersion]:
    storage_object = session.get(StorageObject, storage_object_id)
    if (
        storage_object is None
        or storage_object.status != "active"
        or storage_object.scope != "managed"
    ):
        raise json_error(404, "NOT_FOUND", "Storage object not found.")
    version = session.get(StorageObjectVersion, version_id)
    if version is None or version.storage_object_id != storage_object.id:
        raise json_error(404, "NOT_FOUND", "Storage object version not found.")
    return storage_object, version


def storage_object_and_optional_version(
    session: DatabaseSession,
    storage_object_id: str,
    version_id: str | None,
) -> tuple[StorageObject, StorageObjectVersion | None]:
    storage_object = session.get(StorageObject, storage_object_id)
    if (
        storage_object is None
        or storage_object.status != "active"
        or storage_object.scope != "managed"
    ):
        raise json_error(404, "NOT_FOUND", "Storage object not found.")
    if version_id is None:
        return storage_object, None
    version = session.get(StorageObjectVersion, version_id)
    if version is None or version.storage_object_id != storage_object.id:
        raise json_error(404, "NOT_FOUND", "Storage object version not found.")
    return storage_object, version


@router.get("/objects/{storage_object_id}/llm-digest")
def get_storage_object_llm_digest(
    storage_object_id: str,
    version_id: str | None = Query(default=None),
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    storage_object, version = storage_object_and_optional_version(
        session,
        storage_object_id,
        version_id,
    )
    source = storage_summary_source(storage_object, version, session)
    summary, is_stale, stale_reason = display_file_summary(
        session,
        storage_object_id=storage_object.id,
        storage_object_version_id=(
            source["storage_object_version_id"]
            if isinstance(source["storage_object_version_id"], str)
            else None
        ),
        source_sha256_hex=str(source["sha256_hex"]),
    )
    diff = latest_file_version_diff(
        session,
        storage_object_id=storage_object.id,
        storage_object_version_id=(
            source["storage_object_version_id"]
            if isinstance(source["storage_object_version_id"], str)
            else None
        ),
        source_sha256_hex=str(source["sha256_hex"]),
    )
    return {
        "ok": True,
        "data": {
            "summary": None if summary is None else file_summary_data(summary),
            "source_sha256_hex": source["sha256_hex"],
            "storage_object_version_id": source["storage_object_version_id"],
            "is_stale": is_stale,
            "stale_reason": stale_reason,
            "diff": None
            if diff is None
            else file_version_diff_data(
                diff,
                display_lines=file_version_diff_display_lines(
                    session,
                    storage_object=storage_object,
                    diff=diff,
                ),
            ),
        },
    }


@router.post("/objects/{storage_object_id}/llm-digest")
def prepare_storage_object_llm_digest(
    storage_object_id: str,
    payload: FileSummaryRequest,
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    storage_object, version = storage_object_and_optional_version(
        session,
        storage_object_id,
        payload.storage_object_version_id,
    )
    if not bool(storage_object.llm_input_allowed):
        raise json_error(409, "LLM_INPUT_BLOCKED", "LLM input is blocked for this file.")

    summary = ensure_storage_object_llm_digest(
        session,
        storage_object=storage_object,
        version=version,
        force=True,
    )
    source = storage_summary_source(storage_object, version, session)
    diff = latest_file_version_diff(
        session,
        storage_object_id=storage_object.id,
        storage_object_version_id=(
            source["storage_object_version_id"]
            if isinstance(source["storage_object_version_id"], str)
            else None
        ),
        source_sha256_hex=str(source["sha256_hex"]),
    )
    session.commit()
    return {
        "ok": True,
        "data": {
            "summary": file_summary_data(summary),
            "storage_object": storage_object_data(storage_object, session),
            "source_sha256_hex": source["sha256_hex"],
            "storage_object_version_id": source["storage_object_version_id"],
            "is_stale": False,
            "stale_reason": None,
            "diff": None
            if diff is None
            else file_version_diff_data(
                diff,
                display_lines=file_version_diff_display_lines(
                    session,
                    storage_object=storage_object,
                    diff=diff,
                ),
            ),
        },
    }


@router.delete("/objects/{storage_object_id}/versions/{version_id}/older")
def delete_older_storage_object_versions(
    storage_object_id: str,
    version_id: str,
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    storage_object, version = get_active_storage_object_version(
        session,
        storage_object_id,
        version_id,
    )
    deleted_version_ids = delete_storage_object_versions_up_to(
        storage_object,
        version,
        session=session,
        now=jst_iso(),
    )
    session.commit()
    return {
        "ok": True,
        "data": {
            "storage_object": storage_object_data(storage_object, session),
            "selected_version_id": version.id,
            "deleted_version_ids": deleted_version_ids,
            "deleted_version_count": len(deleted_version_ids),
        },
    }


@router.get("/objects/{storage_object_id}/versions/{version_id}/archive-tree")
def get_storage_object_version_archive_tree(
    storage_object_id: str,
    version_id: str,
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    storage_object, version = get_active_storage_object_version(
        session,
        storage_object_id,
        version_id,
    )
    version_path = storage_object_version_absolute_path(storage_object, version, session)
    if not version_path.is_file():
        raise json_error(404, "NOT_FOUND", "Storage object version file not found.")
    return {
        "ok": True,
        "data": archive_tree_text(version_path),
    }


@router.get("/objects/{storage_object_id}/versions/{version_id}/eml-preview")
def get_storage_object_version_eml_preview(
    storage_object_id: str,
    version_id: str,
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    storage_object, version = get_active_storage_object_version(
        session,
        storage_object_id,
        version_id,
    )
    version_path = storage_object_version_absolute_path(storage_object, version, session)
    if not version_path.is_file():
        raise json_error(404, "NOT_FOUND", "Storage object version file not found.")
    return {"ok": True, "data": eml_preview_data(version_path, session=session)}


@router.get("/objects/{storage_object_id}/versions/{version_id}/content")
def get_storage_object_version_content(
    storage_object_id: str,
    version_id: str,
    session: DatabaseSession = Depends(get_session),
) -> FileResponse:
    storage_object, version = get_active_storage_object_version(
        session,
        storage_object_id,
        version_id,
    )
    version_path = storage_object_version_absolute_path(storage_object, version, session)
    if not version_path.is_file():
        raise json_error(404, "NOT_FOUND", "Storage object version file not found.")
    if should_record_storage_content_view(storage_object):
        record_storage_operation(
            session,
            operation_type="viewed",
            now=jst_iso(),
            storage_object=storage_object,
            details={
                "storage_object_version_id": version.id,
                "version_number": version.version_number,
            },
        )
        session.commit()
    response = FileResponse(
        version_path,
        media_type=version.content_type,
        filename=version.original_filename,
        content_disposition_type="inline",
        headers=storage_content_response_headers(storage_object),
    )
    session.close()
    return response


@router.get("/objects/{storage_object_id}/versions/{version_id}/image-preview")
def get_storage_object_version_image_preview(
    storage_object_id: str,
    version_id: str,
    session: DatabaseSession = Depends(get_session),
) -> Response:
    storage_object, version = get_active_storage_object_version(
        session,
        storage_object_id,
        version_id,
    )
    version_path = storage_object_version_absolute_path(storage_object, version, session)
    if not version_path.is_file():
        raise json_error(404, "NOT_FOUND", "Storage object version file not found.")
    return storage_image_preview_response(
        source_path=version_path,
        content_type=version.content_type,
        filename=version.original_filename,
        sha256_hex=version.sha256_hex,
    )


@router.get("/objects/{storage_object_id}/versions/{version_id}/download")
def download_storage_object_version_content(
    storage_object_id: str,
    version_id: str,
    session: DatabaseSession = Depends(get_session),
) -> FileResponse:
    storage_object, version = get_active_storage_object_version(
        session,
        storage_object_id,
        version_id,
    )
    version_path = storage_object_version_absolute_path(storage_object, version, session)
    if not version_path.is_file():
        raise json_error(404, "NOT_FOUND", "Storage object version file not found.")
    record_storage_operation(
        session,
        operation_type="downloaded",
        now=jst_iso(),
        storage_object=storage_object,
        details={
            "storage_object_version_id": version.id,
            "version_number": version.version_number,
        },
    )
    session.commit()
    response = FileResponse(
        version_path,
        media_type=version.content_type,
        filename=version.original_filename,
        content_disposition_type="attachment",
        headers=NO_STORE_FILE_HEADERS,
    )
    session.close()
    return response


@router.get("/objects/{storage_object_id}")
def get_storage_object(
    storage_object_id: str,
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    storage_object = session.get(StorageObject, storage_object_id)
    if (
        storage_object is None
        or storage_object.status != "active"
        or storage_object.scope != "managed"
    ):
        raise json_error(404, "NOT_FOUND", "Storage object not found.")
    return {
        "ok": True,
        "data": {"storage_object": storage_object_data(storage_object, session)},
    }


@router.get("/objects/{storage_object_id}/archive-tree")
def get_storage_object_archive_tree(
    storage_object_id: str,
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    storage_object = session.get(StorageObject, storage_object_id)
    if storage_object is None or storage_object.status != "active":
        raise json_error(404, "NOT_FOUND", "Storage object not found.")
    object_path = storage_object_absolute_path(storage_object, session)
    if not object_path.is_file():
        raise json_error(404, "NOT_FOUND", "Storage object file not found.")
    return {
        "ok": True,
        "data": archive_tree_text(object_path),
    }


@router.get("/objects/{storage_object_id}/eml-preview")
def get_storage_object_eml_preview(
    storage_object_id: str,
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    storage_object = session.get(StorageObject, storage_object_id)
    if storage_object is None or storage_object.status != "active":
        raise json_error(404, "NOT_FOUND", "Storage object not found.")
    object_path = storage_object_absolute_path(storage_object, session)
    if not object_path.is_file():
        raise json_error(404, "NOT_FOUND", "Storage object file not found.")
    return {"ok": True, "data": eml_preview_data(object_path, session=session)}


@router.get("/objects/{storage_object_id}/content")
def get_storage_object_content(
    storage_object_id: str,
    session: DatabaseSession = Depends(get_session),
) -> FileResponse:
    storage_object = session.get(StorageObject, storage_object_id)
    if storage_object is None or storage_object.status != "active":
        raise json_error(404, "NOT_FOUND", "Storage object not found.")
    object_path = storage_object_absolute_path(storage_object, session)
    if not object_path.is_file():
        raise json_error(404, "NOT_FOUND", "Storage object file not found.")
    if should_record_storage_content_view(storage_object):
        record_storage_operation(
            session,
            operation_type="viewed",
            now=jst_iso(),
            storage_object=storage_object,
        )
        session.commit()
    response = FileResponse(
        object_path,
        media_type=storage_object.content_type,
        filename=storage_object.original_filename,
        content_disposition_type="inline",
        headers=storage_content_response_headers(storage_object),
    )
    session.close()
    return response


@router.get("/objects/{storage_object_id}/image-preview")
def get_storage_object_image_preview(
    storage_object_id: str,
    session: DatabaseSession = Depends(get_session),
) -> Response:
    storage_object = session.get(StorageObject, storage_object_id)
    if storage_object is None or storage_object.status != "active":
        raise json_error(404, "NOT_FOUND", "Storage object not found.")
    object_path = storage_object_absolute_path(storage_object, session)
    if not object_path.is_file():
        raise json_error(404, "NOT_FOUND", "Storage object file not found.")
    if should_record_storage_content_view(storage_object):
        record_storage_operation(
            session,
            operation_type="viewed",
            now=jst_iso(),
            storage_object=storage_object,
        )
        session.commit()
    return storage_image_preview_response(
        source_path=object_path,
        content_type=storage_object.content_type,
        filename=storage_object.original_filename,
        sha256_hex=storage_object.sha256_hex,
    )


@router.get("/objects/{storage_object_id}/download")
def download_storage_object_content(
    storage_object_id: str,
    session: DatabaseSession = Depends(get_session),
) -> FileResponse:
    storage_object = session.get(StorageObject, storage_object_id)
    if storage_object is None or storage_object.status != "active":
        raise json_error(404, "NOT_FOUND", "Storage object not found.")
    object_path = storage_object_absolute_path(storage_object, session)
    if not object_path.is_file():
        raise json_error(404, "NOT_FOUND", "Storage object file not found.")
    record_storage_operation(
        session,
        operation_type="downloaded",
        now=jst_iso(),
        storage_object=storage_object,
    )
    session.commit()
    response = FileResponse(
        object_path,
        media_type=storage_object.content_type,
        filename=storage_object.original_filename,
        content_disposition_type="attachment",
        headers=NO_STORE_FILE_HEADERS,
    )
    session.close()
    return response


@router.post("/tmp")
def upload_temporary_object(
    payload: TemporaryObjectUpload,
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    data = decode_base64_payload(payload.data_base64)
    if len(data) == 0:
        raise json_error(422, "VALIDATION_ERROR", "Temporary object is empty.")
    if len(data) > MAX_TMP_OBJECT_BYTES:
        raise json_error(413, "PAYLOAD_TOO_LARGE", "Temporary object is too large.")

    content_type = (payload.content_type or "application/octet-stream").strip()
    if content_type == "":
        content_type = "application/octet-stream"
    now = jst_iso()
    storage_object = save_storage_object(
        session,
        scope="tmp",
        filename=payload.filename,
        content_type=content_type,
        data=data,
        now=now,
    )
    session.commit()
    return {
        "ok": True,
        "data": {
            "storage_object": storage_object_data(storage_object, session),
        },
    }


@router.post("/contacts/{contact_id}/image")
def upload_contact_image(
    contact_id: str,
    payload: ContactImageUpload,
    session: DatabaseSession = Depends(get_session),
) -> dict[str, object]:
    contact = session.get(Contact, contact_id)
    if contact is None or contact.deleted_at is not None:
        raise json_error(404, "NOT_FOUND", "Contact not found.")
    content_type = payload.content_type.strip().lower()
    if content_type not in CONTACT_IMAGE_CONTENT_TYPES:
        raise json_error(422, "VALIDATION_ERROR", "Unsupported contact image type.")
    data = decode_base64_payload(payload.data_base64)
    if len(data) == 0:
        raise json_error(422, "VALIDATION_ERROR", "Contact image is empty.")
    if len(data) > MAX_CONTACT_IMAGE_BYTES:
        raise json_error(413, "PAYLOAD_TOO_LARGE", "Contact image is too large.")
    content_type, data = resize_contact_image(content_type=content_type, data=data)

    now = jst_iso()
    previous_avatar_url = contact.avatar_url
    storage_object = save_storage_object(
        session,
        scope="contact-images",
        filename=payload.filename,
        content_type=content_type,
        data=data,
        now=now,
    )
    delete_previous_contact_image(
        session,
        avatar_url=previous_avatar_url,
        now=now,
    )
    contact.avatar_url = storage_object_url(storage_object.id)
    contact.updated_at = now
    contact.version += 1
    session.commit()
    return {
        "ok": True,
        "data": {
            "storage_object": storage_object_data(storage_object, session),
            "contact": {
                "id": contact.id,
                "avatar_url": contact.avatar_url,
                "updated_at": contact.updated_at,
                "version": contact.version,
            },
        },
    }
